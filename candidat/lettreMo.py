import io
import json
import logging
import threading
import zipfile

from . import app_messages as messages
from django.db import transaction
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.core.cache import cache as django_cache
from django.views.decorators.http import require_POST

from .decorators import candidat_required
from .models import ModeleLettre

logger = logging.getLogger(__name__)

_MOIS_FR = [
    'janvier','février','mars','avril','mai','juin',
    'juillet','août','septembre','octobre','novembre','décembre',
]

_SALUTATION_MAP = {
    'M.': 'Monsieur', 'M': 'Monsieur',
    'Mme': 'Madame', 'Mme.': 'Madame',
    'Dr.': 'Docteur', 'Pr.': 'Professeur',
}


def _lettre_to_dict_for_render(lettre):
    """Reconstruit le dict plat attendu par lettre_render depuis un objet LettreMotivation."""
    from datetime import date as date_cls
    c  = lettre.candidat
    ip = getattr(c, 'informationPersonnelle', None)
    ct = lettre.contenu

    prenom = c.prenom or (ip.prenom if ip else '') or ''
    nom    = c.nom    or (ip.nom    if ip else '') or ''

    titre_dest = (ct.titreDestinataire if ct else '') or ''
    nom_dest   = (ct.nomDestinataire   if ct else '') or ''
    salutation = _SALUTATION_MAP.get(titre_dest, titre_dest)
    dest_full  = f"{titre_dest} {nom_dest}".strip()

    formule_conge = (ct.formuleConge if ct else '') or ''
    formule_full  = (
        formule_conge
        .replace('[titre]', salutation)
        .replace('[nom]', dest_full)
    )

    date_lettre = (ct.dateLettre if ct else None)
    if date_lettre:
        date_str = f"{date_lettre.day} {_MOIS_FR[date_lettre.month - 1]} {date_lettre.year}"
    else:
        today = date_cls.today()
        date_str = f"{today.day} {_MOIS_FR[today.month - 1]} {today.year}"

    lieu = (ct.lieu if ct else '') or (ip.ville if ip else '') or 'Abidjan'

    return {
        'prenom':            prenom,
        'nom':               nom,
        'adresse':           (ip.adresse    if ip else '') or '',
        'codePostal':        (ip.codePostal if ip else '') or '',
        'ville':             (ip.ville      if ip else '') or '',
        'email':             c.email     or (ip.email     if ip else '') or '',
        'telephone':         c.telephone or (ip.telephone if ip else '') or '',
        'titreDestinataire': titre_dest,
        'nomDestinataire':   nom_dest,
        'posteDestinataire': ct.posteDestinataire.nomPoste if (ct and ct.posteDestinataire) else '',
        'entreprise':        ct.entreprise.nomEntreprise   if (ct and ct.entreprise)        else '',
        'villeEntreprise':   ct.villeEntreprise.nomVille   if (ct and ct.villeEntreprise)   else '',
        'lieu':              lieu,
        'date':              date_str,
        'objet':             (ct.objet   if ct else '') or '',
        'corps':             (ct.corps   if ct else '') or '',
        'formuleConge':      formule_conge,
        'formuleFull':       formule_full,
        'salutation':        salutation,
    }


# ─── Données de démo ──────────────────────────────────────────────────────────

DEMO_LETTRE = {
    'prenom': 'Jean',
    'nom': 'KOUAMÉ',
    'adresse': '12 Rue des Palmiers, Cocody',
    'codePostal': '01 BP 1234',
    'ville': 'Abidjan',
    'pays': "Côte d'Ivoire",
    'email': 'jean.kouame@email.com',
    'telephone': '+225 07 00 00 00 00',
    'titreDestinataire': 'M.',
    'nomDestinataire': 'KONAN',
    'posteDestinataire': 'Directeur des Ressources Humaines',
    'entreprise': "Orange Côte d'Ivoire",
    'villeEntreprise': 'Abidjan',
    'lieu': 'Abidjan',
    'date': '15 mars 2026',
    'objet': 'Candidature au poste de Développeur Full Stack',
    'corps': (
        "Je me permets de vous adresser ma candidature pour le poste de Développeur "
        "Full Stack au sein de votre entreprise, poste que j'ai découvert sur votre site "
        "carrière et qui correspond parfaitement à mon profil et à mes aspirations.\n\n"
        "Titulaire d'un Master en Informatique de l'INP-HB de Yamoussoukro, je dispose "
        "de 3 ans d'expérience en développement web full stack (Django, React, PostgreSQL). "
        "J'ai notamment contribué à des projets d'envergure chez NSIA Technologies, où j'ai "
        "dirigé le développement d'une plateforme bancaire traitant plus de 50 000 transactions "
        "quotidiennes.\n\n"
        "Orange Côte d'Ivoire est une référence incontournable du secteur des télécommunications "
        "en Afrique de l'Ouest. Votre engagement pour la transformation digitale et l'innovation "
        "correspondent parfaitement à mes valeurs et à mes aspirations professionnelles.\n\n"
        "Mon expérience en équipe agile, ma maîtrise des architectures cloud et mon sens du service "
        "me permettraient de contribuer efficacement à vos projets. Je me tiens à votre disposition "
        "pour tout entretien à votre convenance et vous prie de bien vouloir trouver ci-joint mon CV."
    ),
    'formuleFull': "Veuillez agréer, Monsieur KONAN, l'expression de mes salutations distinguées.",
}


# ─── Génération DOCX ──────────────────────────────────────────────────────────

def _generate_docx_lettre(lettre):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    ORANGE = RGBColor(0xF7, 0x7F, 0x00)
    DARK   = RGBColor(0x1a, 0x1a, 0x2e)
    GRAY   = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()

    # ── Mise en page A4 ──────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width   = Cm(21)
    sec.page_height  = Cm(29.7)
    sec.left_margin  = sec.right_margin  = Cm(2.5)
    sec.top_margin   = sec.bottom_margin = Cm(2.2)
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    def para(text='', align=None, bold=False, size=None, color=None,
             space_before=0, space_after=4, italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after  = Pt(space_after)
        if align:
            p.alignment = align
        if text:
            r = p.add_run(text)
            r.bold   = bold
            r.italic = italic
            if size:   r.font.size      = Pt(size)
            if color:  r.font.color.rgb = color
        return p

    # ── Expéditeur (droite) ──────────────────────────────────────────────
    name = f"{lettre.get('prenom', '')} {lettre.get('nom', '')}".strip()
    para(name, align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True, size=13,
         color=DARK, space_after=2)

    contact_lines = []
    if lettre.get('adresse'):
        contact_lines.append(lettre['adresse'])
    if lettre.get('codePostal') or lettre.get('ville'):
        contact_lines.append(
            f"{lettre.get('codePostal', '')} {lettre.get('ville', '')}".strip()
        )
    if lettre.get('telephone'):
        contact_lines.append(lettre['telephone'])
    if lettre.get('email'):
        contact_lines.append(lettre['email'])

    for line in contact_lines:
        para(line, align=WD_ALIGN_PARAGRAPH.RIGHT, size=9, color=GRAY, space_after=1)

    # ── Date & lieu ───────────────────────────────────────────────────────
    date_str = f"{lettre.get('lieu', 'Abidjan')}, le {lettre.get('date', '')}"
    para(date_str, align=WD_ALIGN_PARAGRAPH.RIGHT, size=10,
         color=GRAY, space_before=10, space_after=10)

    # ── Destinataire (gauche) ─────────────────────────────────────────────
    dest_name = f"{lettre.get('titreDestinataire', '')} {lettre.get('nomDestinataire', '')}".strip()
    if dest_name:
        para(dest_name, bold=True, size=11, space_after=1)
    if lettre.get('posteDestinataire'):
        para(lettre['posteDestinataire'], size=10, color=GRAY, space_after=1)
    if lettre.get('entreprise'):
        para(lettre['entreprise'], size=10, space_after=1)
    if lettre.get('villeEntreprise'):
        para(lettre['villeEntreprise'], size=10, space_after=12)

    # ── Objet ─────────────────────────────────────────────────────────────
    if lettre.get('objet'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(16)
        r1 = p.add_run('Objet : ')
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(lettre['objet'])
        r2.bold = True
        r2.font.size = Pt(11)
        r2.font.color.rgb = ORANGE

    # ── Corps ─────────────────────────────────────────────────────────────
    corps = lettre.get('corps', '').strip()
    if corps:
        para(corps, size=11, space_after=10)

    # ── Formule de politesse ──────────────────────────────────────────────
    formule = lettre.get('formuleFull', '').strip()
    if formule:
        para(formule, size=11, space_before=10, space_after=28)

    # ── Signature ─────────────────────────────────────────────────────────
    para(name, bold=True, size=11, color=DARK, align=WD_ALIGN_PARAGRAPH.RIGHT)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Vues ─────────────────────────────────────────────────────────────────────

def modeles_lettre(request):
    from django.core.paginator import Paginator

    cache_key = 'modeles_lettre_data'
    cached = django_cache.get(cache_key)
    if cached is None:
        modeles_qs = ModeleLettre.objects.filter(actif=True).order_by('ordre', 'nom')
        cached = {
            'modeles_list': list(modeles_qs),
            'categories':   ['Tous'] + [label for _, label in ModeleLettre.CATEGORIES],
            'total':        modeles_qs.count(),
            'gratuits':     modeles_qs.filter(premium=False).count(),
            'premium':      modeles_qs.filter(premium=True).count(),
        }
        django_cache.set(cache_key, cached, 1800)

    categories = cached['categories']
    total      = cached['total']
    gratuits   = cached['gratuits']
    premium    = cached['premium']

    PER_PAGE = 20
    page_number = request.GET.get('page', 1)
    paginator = Paginator(cached['modeles_list'], PER_PAGE)
    page_obj  = paginator.get_page(page_number)

    cur   = page_obj.number
    total_pages = paginator.num_pages
    start = max(1, cur - 2)
    end   = min(total_pages, cur + 2)
    if end - start < 4:
        if start == 1:
            end = min(total_pages, start + 4)
        else:
            start = max(1, end - 4)
    page_range = range(start, end + 1)

    params = request.GET.copy()
    params.pop('page', None)
    qs = params.urlencode()

    return render(request, 'candidat/lettreMo/modeles_lettre.html', {
        'modeles':    page_obj,
        'categories': categories,
        'total':      total,
        'gratuits':   gratuits,
        'premium':    premium,
        'page_obj':   page_obj,
        'page_range': page_range,
        'qs':         qs,
    })


def apercu_lettre(request, template_id):
    """Prévisualisation du modèle (données de démo, sans connexion requise)."""
    modele = ModeleLettre.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return redirect('candidat:modeles_lettre')
    return render(request, f'candidat/lettreMo/modeles/{modele.slug}.html', {
        'lettre': DEMO_LETTRE,
        'modele': modele,
        'apercu': True,
    })


@candidat_required
def creer_lettre(request, template_id):
    from .models import LettreMotivation
    modele = ModeleLettre.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        messages.error(request, "Ce modèle est introuvable ou n'est plus disponible.")
        return redirect('candidat:modeles_lettre')
    if modele.premium:
        messages.error(request, 'Ce modèle est réservé aux membres Premium.')
        return redirect('candidat:modeles_lettre')

    modeles = ModeleLettre.objects.filter(actif=True).order_by('ordre', 'nom')

    modeles_json = json.dumps([
        {
            'id':      m.id,
            'nom':     m.nom,
            'slug':    m.slug,
            'couleur': m.couleur,
            'premium': m.premium,
            'famille': m.famille,
            'apercu':  m.apercu.url if m.apercu else None,
        }
        for m in modeles
    ], ensure_ascii=False)

    c  = request.candidat
    ip = getattr(c, 'informationPersonnelle', None)

    def _f(direct, legacy=''):
        return direct or legacy or ''

    candidat_info = {
        'prenom':     _f(c.prenom,    ip.prenom     if ip else ''),
        'nom':        _f(c.nom,       ip.nom        if ip else ''),
        'email':      _f(c.email,     ip.email      if ip else ''),
        'telephone':  _f(c.telephone, ip.telephone  if ip else ''),
        'adresse':    _f(c.adresse,   ip.adresse    if ip else ''),
        'codePostal': _f('',          ip.codePostal if ip else ''),
        'ville':      _f('',          ip.ville      if ip else ''),
    }

    # Chargement depuis la BD si on modifie une lettre existante
    lettre_initial = None
    lettre_id_param = request.GET.get('lettre_id')
    if lettre_id_param:
        lettre = (
            LettreMotivation.objects
            .filter(pk=lettre_id_param, candidat=c, archive=False)
            .select_related('contenu', 'contenu__entreprise',
                            'contenu__posteDestinataire', 'contenu__pays', 'contenu__villeEntreprise')
            .first()
        )
        if lettre:
            ct = lettre.contenu
            # Dict Python — json_script sérialise une seule fois (pas de double-encodage)
            lettre_initial = {
                'lettreId':          lettre.pk,
                'nomLettre':         lettre.nomLettre,
                'titreDestinataire': ct.titreDestinataire if ct else '',
                'nomDestinataire':   ct.nomDestinataire   if ct else '',
                'posteDestinataire': ct.posteDestinataire.nomPoste if ct and ct.posteDestinataire else '',
                'posteDestId':       ct.posteDestinataire.pk        if ct and ct.posteDestinataire else None,
                'entreprise':        ct.entreprise.nomEntreprise    if ct and ct.entreprise else '',
                'entrepriseId':      ct.entreprise.pk               if ct and ct.entreprise else None,
                'paysNom':           ct.pays.nomPays                if ct and ct.pays else '',
                'paysId':            ct.pays.pk                     if ct and ct.pays else None,
                'villeEntreprise':   ct.villeEntreprise.nomVille    if ct and ct.villeEntreprise else '',
                'villeEntrepriseId': ct.villeEntreprise.pk          if ct and ct.villeEntreprise else None,
                'lieu':              ct.lieu                        if ct else '',
                'dateLettre':        ct.dateLettre.isoformat()      if ct and ct.dateLettre else '',
                'objet':             ct.objet                       if ct else '',
                'corps':             ct.corps                       if ct else '',
                'formuleConge':      ct.formuleConge                if ct else '',
            }

    return render(request, 'candidat/lettreMo/creer_lettre.html', {
        'modele':         modele,
        'modeles':        modeles,
        'modeles_json':   modeles_json,
        'candidat_info':  candidat_info,
        'lettre_initial': lettre_initial,
    })


@candidat_required
def modifier_lettre(request, lettre_id):
    """Redirige vers l'éditeur de la lettre existante."""
    from .models import LettreMotivation
    lettre = (
        LettreMotivation.objects
        .filter(pk=lettre_id, candidat=request.candidat, archive=False)
        .select_related('modele')
        .first()
    )
    if not lettre or not lettre.modele:
        messages.error(request, 'Lettre introuvable ou modèle indisponible.')
        return redirect('candidat:profil')
    from django.urls import reverse
    return redirect(
        reverse('candidat:creer_lettre', args=[lettre.modele.pk]) + f'?lettre_id={lettre.pk}'
    )


@candidat_required
def api_lister_lettres(request):
    """JSON : liste des lettres non archivées du candidat."""
    from django.urls import reverse
    from django.utils import timezone
    from datetime import timedelta

    lettres = (
        request.candidat.lettres.filter(archive=False)
        .select_related('modele', 'contenu', 'contenu__entreprise')
        .prefetch_related('photos')
        .order_by('-dateModification')
    )
    pending_threshold = timezone.now() - timedelta(seconds=120)

    def _serialize(l):
        ct     = l.contenu
        photos = list(l.photos.all())
        mini   = photos[0] if photos else None
        thumb  = mini.image.url if (mini and mini.image) else ''
        has_pdf = bool(l.lettrePdf)
        pdf_pending = (not has_pdf or not thumb) and l.dateModification > pending_threshold
        modele_actif = l.modele and l.modele.actif
        return {
            'id':                l.pk,
            'nomLettre':         l.nomLettre or '',
            'modele_nom':        l.modele.nom if l.modele else '',
            'modele_indispo':    not modele_actif,
            'objet':             ct.objet if ct else '',
            'entreprise':        ct.entreprise.nomEntreprise if ct and ct.entreprise else '',
            'date_modification': l.dateModification.strftime('%d/%m/%Y à %H:%M'),
            'modifier_url':      reverse('candidat:modifier_lettre', args=[l.pk]) if modele_actif else '',
            'pdf_url':           l.lettrePdf.url if has_pdf else '',
            'thumbnail_url':     thumb,
            'pdf_pending':       pdf_pending,
            'images_url':        reverse('candidat:api_images_lettre', args=[l.pk]) if thumb else '',
        }

    return JsonResponse({'ok': True, 'lettres': [_serialize(l) for l in lettres]})


@candidat_required
@require_POST
def archiver_lettre(request, lettre_id):
    """Archive (soft-delete) une lettre."""
    from .models import LettreMotivation
    lettre = LettreMotivation.objects.filter(pk=lettre_id, candidat=request.candidat).first()
    if not lettre:
        return JsonResponse({'ok': False, 'message': 'Lettre introuvable.'}, status=404)
    lettre.archiver()
    return JsonResponse({'ok': True, 'message': 'Lettre archivée.'})


@candidat_required
@require_POST
def regenerer_lettre_artefacts(request, lettre_id):
    """Régénère le PDF et la miniature d'une lettre sauvegardée."""
    from .models import LettreMotivation
    lettre = (
        LettreMotivation.objects
        .filter(pk=lettre_id, candidat=request.candidat, archive=False)
        .select_related('modele', 'contenu', 'contenu__entreprise',
                        'contenu__posteDestinataire', 'contenu__villeEntreprise')
        .first()
    )
    if not lettre:
        return JsonResponse({'ok': False, 'message': 'Lettre introuvable.'}, status=404)
    if not lettre.modele or not lettre.modele.actif:
        return JsonResponse(
            {'ok': False, 'message': 'Modèle indisponible — impossible de régénérer.'},
            status=400,
        )
    try:
        ok = lettre.generer_artefacts(request)
    except Exception as exc:
        logger.exception('Échec régénération lettre %s', lettre_id)
        return JsonResponse(
            {'ok': False, 'message': f'Erreur : {exc}'},
            status=500,
        )
    if not ok:
        return JsonResponse(
            {'ok': False, 'message': 'Régénération échouée (modèle ou contenu incomplet).'},
            status=500,
        )
    photo = lettre.photos.order_by('numeroPage').first()
    return JsonResponse({
        'ok':            True,
        'message':       'Lettre mise à jour — PDF et image régénérés.',
        'pdf_url':       lettre.lettrePdf.url if lettre.lettrePdf else '',
        'thumbnail_url': photo.image.url if (photo and photo.image) else '',
    })


@candidat_required
def api_images_lettre(request, lettre_id):
    """Télécharge la miniature PNG de la lettre."""
    from .models import LettreMotivation
    lettre = LettreMotivation.objects.filter(
        pk=lettre_id, candidat=request.candidat, archive=False
    ).first()
    if not lettre:
        return HttpResponse('Lettre introuvable.', status=404)
    photo = lettre.photos.order_by('numeroPage').first()
    if not photo:
        return HttpResponse('Aucune image disponible.', status=404)
    nom = lettre.nomLettre or f'Lettre_{lettre.pk}'
    nom = ''.join(c for c in nom if c not in r'\/:*?"<>|')
    response = HttpResponse(photo.image.read(), content_type='image/png')
    response['Content-Disposition'] = f'attachment; filename="{nom}.png"'
    return response


@candidat_required
@require_POST
def telecharger_lettre(request, template_id, fmt):
    """Génère et télécharge la lettre en PDF, PNG, JPG ou DOCX."""
    if fmt not in ('pdf', 'png', 'jpg', 'docx'):
        return HttpResponse('Format non supporté', status=400)

    try:
        lettre_data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return HttpResponse('Données invalides', status=400)

    modele = ModeleLettre.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return HttpResponse('Modèle introuvable', status=404)

    prenom   = lettre_data.get('prenom', 'Ma').strip().replace(' ', '-')
    nom      = lettre_data.get('nom', 'Lettre').strip().replace(' ', '-')
    filename = f"Lettre-{prenom}-{nom}"

    from . import lettre_render

    try:
        if fmt == 'pdf':
            data         = lettre_render.render_pdf(modele, lettre_data, request)
            content_type = 'application/pdf'
            dl_filename  = f"{filename}.pdf"

        elif fmt in ('png', 'jpg'):
            data         = lettre_render.render_image(modele, lettre_data, request, fmt=fmt)
            content_type = 'image/jpeg' if fmt == 'jpg' else 'image/png'
            dl_filename  = f"{filename}.{fmt}"

        else:  # docx
            data         = _generate_docx_lettre(lettre_data)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            dl_filename  = f"{filename}.docx"

    except ImportError as e:
        lib = str(e).split("'")[-2] if "'" in str(e) else str(e)
        return HttpResponse(
            f'Bibliothèque manquante : {lib}\n'
            f'Installez-la avec : pip install playwright pymupdf pillow python-docx\n'
            f'Puis : playwright install chromium',
            status=500, content_type='text/plain; charset=utf-8',
        )
    except Exception as e:
        return HttpResponse(
            f'Erreur de génération : {e}',
            status=500, content_type='text/plain; charset=utf-8',
        )

    resp = HttpResponse(data, content_type=content_type)
    resp['Content-Disposition'] = f'attachment; filename="{dl_filename}"'
    resp['Content-Length'] = len(data)
    return resp


def _schedule_lettre_artefacts(lettre_id, request):
    """Lance la génération PDF + miniature dans un thread daemon.

    Identique au pattern de `_schedule_artefacts_generation` pour les CV :
    retourne immédiatement → la réponse HTTP n'attend pas Playwright (~2-5s).
    Le thread re-fetche la lettre depuis la DB et appelle generer_artefacts().
    Erreurs : loggées, jamais propagées.
    """
    def _run():
        from django.db import close_old_connections
        from .models import LettreMotivation
        try:
            close_old_connections()
            lettre = (
                LettreMotivation.objects
                .filter(pk=lettre_id)
                .select_related('modele', 'contenu', 'contenu__entreprise',
                                'contenu__posteDestinataire', 'contenu__villeEntreprise',
                                'candidat', 'candidat__informationPersonnelle')
                .first()
            )
            if lettre:
                lettre.generer_artefacts(request)
        except Exception:
            logger.exception('Génération async artefacts lettre %s échouée', lettre_id)
        finally:
            close_old_connections()

    threading.Thread(
        target=_run,
        name=f'lettre-artefacts-{lettre_id}',
        daemon=True,
    ).start()


@candidat_required
@require_POST
def sauvegarder_lettre(request, template_id):
    """Sauvegarde ou met à jour une LettreMotivation en base de données."""
    from datetime import date as date_type
    from .models import LettreMotivation, LettreContenu
    from referentiel.models import RaisonSociale, Poste, Pays, Ville

    try:
        data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'ok': False, 'error': 'Données invalides'}, status=400)

    modele = ModeleLettre.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return JsonResponse({'ok': False, 'error': 'Modèle introuvable'}, status=404)

    c  = request.candidat
    ip = getattr(c, 'informationPersonnelle', None)

    # ── Résolution FK entreprise ──────────────────────────────────────────────
    entreprise = None
    nom_ent = (data.get('entreprise') or '').strip()
    if nom_ent:
        ent_id = data.get('entrepriseId')
        if ent_id:
            entreprise = RaisonSociale.objects.filter(pk=ent_id).first()
        if not entreprise:
            entreprise, _ = RaisonSociale.objects.get_or_create(nomEntreprise=nom_ent)

    # ── Résolution FK poste destinataire ─────────────────────────────────────
    poste_dest = None
    nom_poste = (data.get('posteDestinataire') or '').strip()
    if nom_poste:
        poste_id = data.get('posteDestId')
        if poste_id:
            poste_dest = Poste.objects.filter(pk=poste_id).first()
        if not poste_dest:
            poste_dest, _ = Poste.objects.get_or_create(nomPoste=nom_poste)

    # ── Résolution FK pays ───────────────────────────────────────────────────
    pays_obj = None
    nom_pays = (data.get('paysNom') or '').strip()
    if nom_pays:
        pays_id = data.get('paysId')
        if pays_id:
            pays_obj = Pays.objects.filter(pk=pays_id).first()
        if not pays_obj:
            pays_obj = Pays.objects.filter(nomPays__iexact=nom_pays).first()
            if not pays_obj:
                pays_obj = Pays.objects.create(nomPays=nom_pays)

    # ── Résolution FK ville entreprise ───────────────────────────────────────
    ville_ent = None
    nom_ville = (data.get('villeEntreprise') or '').strip()
    if nom_ville:
        ville_id = data.get('villeEntrepriseId')
        if ville_id:
            ville_ent = Ville.objects.filter(pk=ville_id).first()
        if not ville_ent:
            qs = Ville.objects.filter(nomVille__iexact=nom_ville)
            if pays_obj:
                qs = qs.filter(pays=pays_obj)
            ville_ent = qs.first()
            if not ville_ent:
                ville_ent = Ville.objects.create(nomVille=nom_ville, pays=pays_obj)

    # ── Date ──────────────────────────────────────────────────────────────────
    date_lettre = None
    date_str = (data.get('dateLettre') or '').strip()
    if date_str:
        try:
            date_lettre = date_type.fromisoformat(date_str)
        except ValueError:
            pass

    # ── Nom de la lettre ──────────────────────────────────────────────────────
    nom_lettre = (data.get('nomLettre') or '').strip()
    if not nom_lettre:
        nom_lettre = f"Lettre — {nom_ent or 'Sans titre'}"

    # ── Récupération ou création ───────────────────────────────────────────────
    lettre_id = data.get('lettreId')
    lettre = None
    if lettre_id:
        lettre = LettreMotivation.objects.filter(pk=lettre_id, candidat=c).select_related('contenu').first()

    contenu = (lettre.contenu if lettre and lettre.contenu else None) or LettreContenu()
    contenu.informationPersonnelle = ip
    contenu.entreprise             = entreprise
    contenu.titreDestinataire      = (data.get('titreDestinataire') or '').strip()
    contenu.nomDestinataire        = (data.get('nomDestinataire')   or '').strip()
    contenu.posteDestinataire      = poste_dest
    contenu.pays                   = pays_obj
    contenu.villeEntreprise        = ville_ent
    contenu.dateLettre             = date_lettre
    contenu.lieu                   = (data.get('lieu')         or '').strip()
    contenu.objet                  = (data.get('objet')        or '').strip()
    contenu.corps                  = (data.get('corps')        or '').strip()
    contenu.formuleConge           = (data.get('formuleConge') or '').strip()
    with transaction.atomic():
        contenu.save()

        if not lettre:
            lettre = LettreMotivation(candidat=c)
        lettre.modele    = modele
        lettre.contenu   = contenu
        lettre.nomLettre = nom_lettre
        lettre.save()

    # Génère PDF + miniature dans un thread daemon → réponse HTTP immédiate
    _schedule_lettre_artefacts(lettre.pk, request)

    return JsonResponse({
        'ok':         True,
        'lettreId':   lettre.pk,
        'nomLettre':  lettre.nomLettre,
        'pdf_url':    lettre.lettrePdf.url if lettre.lettrePdf else '',
        'pdf_pending': True,
    })
