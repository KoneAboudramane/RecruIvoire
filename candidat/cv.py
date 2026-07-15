import io
import json
import logging
import threading
import zipfile

logger = logging.getLogger(__name__)

from . import app_messages as messages
from django.core.files.base import ContentFile
from django.db import transaction
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect
from django.template.loader import render_to_string
from django.urls import reverse
from django.core.cache import cache as django_cache
from django.views.decorators.http import require_POST

from .decorators import candidat_required
from .models import (
    CV, CVContenu, InformationPersonnelle, ModeleCV, PhotoCV,
    Formation, ExperienceProfessionnelle, PosteOccupe,
    Competence, CandidatLangue, CentreInteret,
)


# ─── Helpers ──────────────────────────────────────────────────────────────────

_MOIS = ['Jan.', 'Fév.', 'Mar.', 'Avr.', 'Mai', 'Juin',
         'Juil.', 'Août', 'Sep.', 'Oct.', 'Nov.', 'Déc.']

def _build_niveau_index():
    """Charge `referentiel.Niveau` en mémoire : {id: Niveau}."""
    from referentiel.models import Niveau
    return {n.id: n for n in Niveau.objects.all()}


def _niveau_etoiles_label(val, niveaux_idx, type_):
    """`val` = id Niveau (nouveau format) | int 1-5 stars legacy (compétence)
    | code CEFR legacy (langue) | None.
    Retourne (nbEtoiles 0-5, libellé affichable)."""
    if val in (None, ''):
        return (0, '')
    # Format nouveau : id Niveau (int ou str numérique)
    try:
        ival = int(val)
        if ival in niveaux_idx:
            n = niveaux_idx[ival]
            return (n.nbEtoiles, n.libelle or n.nomNiveau)
        # Legacy compétence : 1-5 étoiles directes
        if type_ == 'competence' and 1 <= ival <= 5:
            for n in niveaux_idx.values():
                if n.type == 'competence' and n.nbEtoiles == ival:
                    return (n.nbEtoiles, n.libelle or n.nomNiveau)
    except (TypeError, ValueError):
        pass
    # Legacy langue : code CEFR
    if type_ == 'langue' and isinstance(val, str):
        for n in niveaux_idx.values():
            if n.type == 'langue' and n.nomNiveau == val:
                return (n.nbEtoiles, n.libelle or n.nomNiveau)
    return (0, '')


def _fmt_month(s):
    if not s:
        return ''
    s = str(s)
    if '-' not in s:
        return s  # année seule
    try:
        y, m = s.split('-')
        return _MOIS[int(m) - 1] + ' ' + y
    except Exception:
        return s


def _fmt_periode(debut, fin, en_cours):
    d = _fmt_month(debut)
    f = 'Présent' if en_cours else _fmt_month(fin)
    if not d and not f:
        return ''
    if not d:
        return f
    if not f and not en_cours:
        return d
    return f'{d} — {f}'


def _exp_visible(e):
    return bool(e.get('entreprise') or any(p.get('titre') for p in e.get('postes', [])))


def _exp_global_periode(e):
    postes = e.get('postes', [])
    active = [p for p in postes if p.get('debut') or p.get('enCours')]
    if not active:
        return ''
    en_cours = any(p.get('enCours') for p in active)
    debuts   = sorted(p['debut'] for p in active if p.get('debut'))
    fins     = sorted(p['fin'] for p in active if not p.get('enCours') and p.get('fin'))
    return _fmt_periode(debuts[0] if debuts else '', fins[-1] if fins else '', en_cours)


def _exp_titre(e):
    postes = e.get('postes', [])
    if len(postes) > 1:
        return e.get('entreprise') or '—'
    return (postes[0].get('titre') if postes else '') or '—'


def _exp_periode(e):
    # Priorité aux dates au niveau expérience (nouvelle architecture)
    if e.get('debut') or e.get('enCours'):
        return _fmt_periode(e.get('debut', ''), e.get('fin', ''), e.get('enCours', False))
    # Rétrocompatibilité : calcul depuis les postes (anciennes données)
    return _exp_global_periode(e)


def _preprocess_cv(cv):
    """Ajoute les champs pré-calculés (periode, niveau_pct, dots, listes filtrées)."""
    for e in cv.get('experiences', []):
        if 'postes' not in e:
            e['postes'] = [{
                'titre':   e.pop('poste', '') or '',
                'debut':   e.pop('debut', '') or '',
                'fin':     e.pop('fin', '')   or '',
                'enCours': bool(e.pop('enCours', False)),
            }]
        for p in e.get('postes', []):
            p['periode'] = _fmt_periode(p.get('debut', ''), p.get('fin', ''), p.get('enCours', False))
        e['periode'] = _exp_periode(e)
        e['titre']   = _exp_titre(e)
        for mc in e.get('missionsClient', []):
            mc['periode'] = _fmt_periode(mc.get('debut', ''), mc.get('fin', ''), mc.get('enCours', False))
        e['missionsClient_ok'] = [mc for mc in e.get('missionsClient', []) if mc.get('client')]
    for f in cv.get('formations', []):
        f['periode'] = _fmt_periode(f.get('debut', ''), f.get('fin', ''), f.get('enCours', False))
        # Pour les certifications : période = date obtention / expiration
        if f.get('typeSortie') == 'certification':
            obtenu = _fmt_month(f.get('fin', ''))
            expire = _fmt_month(f.get('expiration', ''))
            if obtenu and expire:
                f['periode_certif'] = f'Obtenu {obtenu} · Exp. {expire}'
            elif obtenu:
                f['periode_certif'] = f'Obtenu {obtenu}'
            elif expire:
                f['periode_certif'] = f'Exp. {expire}'
            else:
                f['periode_certif'] = ''
        else:
            f['periode_certif'] = f['periode']
    for b in cv.get('benevs', []):
        b['periode'] = _fmt_periode(b.get('debut', ''), b.get('fin', ''), b.get('enCours', False))
    niveaux_idx = _build_niveau_index()
    for c in cv.get('competences', []):
        nb, lbl = _niveau_etoiles_label(c.get('niveau'), niveaux_idx, 'competence')
        c['niveau_nb']    = nb
        c['niveau_label'] = lbl
        c['niveau_pct']   = nb * 20
    for l in cv.get('langues', []):
        nb, lbl = _niveau_etoiles_label(l.get('niveau'), niveaux_idx, 'langue')
        l['niveau_nb']    = nb
        l['niveau_label'] = lbl
        l['dots']         = [i <= nb for i in range(1, 6)]

    # Filtre `visible !== false` : un item masqué dans l'éditeur est exclu
    # du rendu (CV téléchargé). Les items sans champ `visible` sont visibles
    # par défaut (rétrocompat avec données pré-bascule).
    def _vis(it):
        return it.get('visible') is not False if isinstance(it, dict) else True

    cv['competences_ok']  = [c for c in cv.get('competences', []) if _vis(c) and c.get('nom')]
    cv['langues_ok']      = [l for l in cv.get('langues', [])     if _vis(l) and l.get('nom')]
    cv['interets_ok']     = [i for i in cv.get('interets', [])    if i]
    cv['experiences_ok']  = [e for e in cv.get('experiences', []) if _vis(e) and _exp_visible(e)]
    cv['formations_ok']   = [f for f in cv.get('formations', [])  if _vis(f) and (f.get('diplome') or f.get('ecole'))]
    cv['certifs_ok']      = [c for c in cv.get('certifs', [])     if c.get('nom')]
    cv['projets_ok']      = [p for p in cv.get('projets', [])     if _vis(p) and p.get('nom')]
    cv['benevs_ok']       = [b for b in cv.get('benevs', [])      if _vis(b) and b.get('orga')]
    # Sous-listes séparées pour le rendu DOCX
    cv['diplomes_ok']     = [f for f in cv['formations_ok'] if f.get('typeSortie', 'diplome') != 'certification']
    cv['certifs_form_ok'] = [f for f in cv['formations_ok'] if f.get('typeSortie') == 'certification']
    return cv


# ─── Snapshot du CV (figé au moment du save) ─────────────────────────────────

# Sections rubrique → clé du snapshot. Sert au build du snapshot et au calcul
# des suggestions (diff snapshot ↔ candidat live).
_SNAPSHOT_SECTIONS = (
    'experiences', 'formations', 'competences', 'langues',
    'interets', 'projets', 'benevs',
)

# Champs identité / contact / toggles à conserver dans le snapshot.
# Photo volontairement exclue : c'est une URL/base64 attachée au candidat,
# rechargée depuis _photo_url() à la lecture (évite de figer du base64 lourd).
_SNAPSHOT_META_KEYS = (
    'prenom', 'nom', 'titre', 'email', 'telephone',
    'ville', 'pays', 'adresse', 'age', 'linkedin', 'portfolio',
    'permis', 'profil',
    'showCertif', 'showProjets', 'showBenev', 'showRef',
    'interetsMasques',
)


def _build_snapshot(cv_data, elements_masques):
    """Construit le dict figé à stocker dans `CVContenu.donneesSnapshot`.

    Inclut :
      - identité / contact / profil (snapshot des champs candidat au moment du save)
      - toutes les rubriques (avec leurs `visible` / IDs)
      - les toggles d'affichage
      - la carte de masquage (`elementsMasques`) calculée plus haut
    Volontairement minimal et plat : on stocke ce que l'éditeur consomme.
    """
    snap = {k: cv_data.get(k) for k in _SNAPSHOT_META_KEYS if k in cv_data}
    for section in _SNAPSHOT_SECTIONS:
        snap[section] = cv_data.get(section) or []
    snap['elementsMasques'] = elements_masques or {}
    return snap


# ─── Reconstruction dict éditeur depuis l'ORM ────────────────────────────────

def _ym(d):
    """date → 'YYYY-MM' (format attendu par les <input type='month'>)."""
    if not d:
        return ''
    if isinstance(d, str):
        return d[:7] if len(d) >= 7 and d[4:5] == '-' else d
    try:
        return d.strftime('%Y-%m')
    except Exception:
        return ''


def _calculer_age(date_naissance):
    if not date_naissance:
        return ''
    try:
        from datetime import date as _date
        today = _date.today()
        years = today.year - date_naissance.year - (
            (today.month, today.day) < (date_naissance.month, date_naissance.day)
        )
        return str(years) if years > 0 else ''
    except Exception:
        return ''


def _photo_url(candidat):
    photo = getattr(candidat, 'photoProfil', None)
    if photo and getattr(photo, 'name', ''):
        try:
            return photo.url
        except Exception:
            return ''
    return ''


def _cv_photo_url(cv_obj):
    """URL de la photo propre au CV (snapshot au save), fallback candidat."""
    photo = getattr(cv_obj, 'photoProfil', None)
    if photo and getattr(photo, 'name', ''):
        try:
            return photo.url
        except Exception:
            pass
    return _photo_url(cv_obj.candidat)


def _permis_text(candidat, info):
    """Texte permis : M2M typePermis du candidat → libellés concaténés ; sinon legacy info.permis."""
    try:
        permis = list(candidat.typePermis.values_list('nomPermis', flat=True))
        if permis:
            return ', '.join(permis)
    except Exception:
        pass
    return _txt(getattr(info, 'permis', '')) if info else ''


def _experience_to_dict(e):
    ville = e.ville or ''
    pays  = e.paysLibre or (e.pays.nomPays if e.pays_id else '')
    return {
        'id'         : e.id,
        'entreprise' : e.entrepriseLibre or (e.entreprise.nomEntreprise if e.entreprise_id else ''),
        'ville'      : ville,
        'pays'       : pays,
        'lieu'       : ', '.join(p for p in (ville, pays) if p),
        'debut'      : _ym(e.dateDebut),
        'fin'        : _ym(e.dateFin),
        'enCours'    : bool(e.enCours),
        'postes'     : [{
            'id'     : p.id,
            'titre'  : p.titreLibre or (p.poste.nomPoste if p.poste_id else ''),
            'debut'  : _ym(p.dateDebut),
            'fin'    : _ym(p.dateFin),
            'enCours': bool(p.enCours),
        } for p in e.postes.all()],
        'hasMissionsClient': e.missionsClient.exists(),
        'missionsClient'   : [{
            'id'     : mc.id,
            'client' : mc.clientLibre or (mc.client.nomEntreprise if mc.client_id else ''),
            'pays'   : mc.paysLibre or (mc.pays.nomPays if mc.pays_id else ''),
            'ville'  : mc.ville or '',
            'debut'  : _ym(mc.dateDebut),
            'fin'    : _ym(mc.dateFin),
            'enCours': bool(mc.enCours),
            'desc'   : mc.description or '',
        } for mc in e.missionsClient.all()],
    }


def _formation_to_dict(f):
    return {
        'id'         : f.id,
        'typeSortie' : f.typeSortie or 'diplome',
        'diplome'    : f.diplomeLibre or (f.diplomeRef.nomDiplome if f.diplomeRef_id else ''),
        'domaine'    : f.domaineLibre or (f.domaine.nomDomaine if f.domaine_id else ''),
        'niveauEtude': f.niveauEtude.nomNiveau if f.niveauEtude_id else '',
        'ecole'      : f.ecoleLibre or (f.institution.nomInstitution if f.institution_id else ''),
        'lieu'       : f.ville or '',
        'debut'      : _ym(f.dateDebut),
        'fin'        : _ym(f.dateFin),
        'enCours'    : bool(f.enCours),
        'desc'       : f.description or '',
        'numero'     : f.numero or '',
        'expiration' : _ym(f.expiration),
    }


def _cv_to_dict(cv_obj):
    """Reconstruit le dict consommé par l'éditeur Alpine.js et `_preprocess_cv`.

    Priorité au snapshot figé : si `CVContenu.donneesSnapshot` est présent, on
    le retourne tel quel — le CV apparaît exactement comme au moment du save.
    `suggestions` est joint au dict : items présents chez le candidat (table
    live) et absents du snapshot, que l'éditeur affiche dans un panneau « à
    ajouter au CV ». Le candidat peut accepter (item bascule dans la rubrique)
    ou ignorer (suggestion supprimée pour ce CV).

    Fallback (CV legacy sans snapshot) : reconstruction depuis l'ORM (M2M de
    `CVContenu` + FK info perso, fallback live candidat). Comportement
    historique préservé pour les CV créés avant le rollout.
    """
    contenu = cv_obj.contenu
    snapshot = (contenu.donneesSnapshot if contenu else None) or {}
    if snapshot:
        return _from_snapshot(cv_obj, snapshot)
    return _from_orm(cv_obj)


def _from_snapshot(cv_obj, snapshot):
    """Retourne le dict éditeur depuis un CV figé + ses suggestions."""
    candidat = cv_obj.candidat
    data = dict(snapshot)

    # Champs portés par CV (et pas par le snapshot) : titre, profil → on prend
    # la valeur de CV qui peut avoir été mise à jour côté admin/ORM.
    data['nomCv']  = cv_obj.nomCv  or data.get('nomCv',  '')
    data['titre']  = cv_obj.titre  or data.get('titre',  '') or candidat.titreProfessionnel or ''
    data['profil'] = cv_obj.profil or data.get('profil', '') or candidat.profilCV or ''

    # Photo : on relit la photo PROPRE au CV (snapshot au save) ; fallback
    # vers la photo de profil candidat si le CV n'en a pas (CVs legacy).
    data['photo'] = _cv_photo_url(cv_obj)

    data['suggestions'] = _compute_suggestions(candidat, snapshot)
    return data


def _from_orm(cv_obj):
    """Reconstruction historique (CV pré-snapshot) — comportement préservé.

    Identité depuis le snapshot `CVContenu.informationPersonnelle`, rubriques
    depuis les M2M de `CVContenu`, fallback live candidat si M2M vide.
    """
    candidat = cv_obj.candidat
    contenu  = cv_obj.contenu
    info     = contenu.informationPersonnelle if contenu else None
    masques  = (contenu.elementsMasques if contenu else None) or {}

    def _apply_vis(items, section):
        """Marque visible:false les items dont l'id DB est dans masques[section]."""
        masked_ids = set(masques.get(section) or [])
        for item in items:
            item_id = item.get('id')
            item['visible'] = item_id not in masked_ids
        return items

    # ── Identité (snapshot info perso prioritaire, fallback candidat) ──────
    nom       = _txt(getattr(info, 'nom',       '') or candidat.nom)
    prenom    = _txt(getattr(info, 'prenom',    '') or candidat.prenom)
    email     = _txt(getattr(info, 'email',     '') or candidat.email)
    telephone = _txt(getattr(info, 'telephone', '') or candidat.telephone)
    adresse   = _txt(getattr(info, 'adresse',   '') or candidat.adresse)
    ville     = _txt(getattr(info, 'ville',     ''))
    pays      = _txt(getattr(info, 'pays',      ''))
    permis    = _permis_text(candidat, info)

    # ── Collections : M2M (figé), fallback candidat live si M2M vide ──────
    def _m2m_or_live(m2m_attr, live_qs, mapper):
        items = list(m2m_attr.all()) if contenu else []
        if not items:
            items = list(live_qs)
        return [mapper(x) for x in items] if mapper else items

    if contenu:
        experiences_qs = contenu.experiences.all().prefetch_related('postes', 'missionsClient')
    else:
        experiences_qs = []
    experiences = [_experience_to_dict(e) for e in experiences_qs]
    if not experiences:
        live = candidat.experiencesProfessionnelles.all().prefetch_related('postes', 'missionsClient')
        experiences = [_experience_to_dict(e) for e in live]

    formations = [_formation_to_dict(f) for f in (contenu.formations.all() if contenu else [])]
    if not formations:
        formations = [_formation_to_dict(f) for f in candidat.formations.all()]

    from .niveau_resolver import resolve_competence_niveau_id, resolve_langue_niveau_id
    niveaux_idx = _build_niveau_index()

    competences = [{
        'id'    : c.id,
        'nom'   : c.nomLibre or (c.typeCompetence.nomCompetence if c.typeCompetence_id else ''),
        'niveau': resolve_competence_niveau_id(c, niveaux_idx),
    } for c in (contenu.competences.all() if contenu else [])]
    if not competences:
        competences = [{
            'id'    : c.id,
            'nom'   : c.nomLibre or (c.typeCompetence.nomCompetence if c.typeCompetence_id else ''),
            'niveau': resolve_competence_niveau_id(c, niveaux_idx),
        } for c in candidat.competences.all()]

    langues = [{
        'id'    : l.id,
        'nom'   : l.nomLibre or (l.langue.nomLangue if l.langue_id else ''),
        'niveau': resolve_langue_niveau_id(l, niveaux_idx),
    } for l in (contenu.langues.all() if contenu else [])]
    if not langues:
        langues = [{
            'id'    : l.id,
            'nom'   : l.nomLibre or (l.langue.nomLangue if l.langue_id else ''),
            'niveau': resolve_langue_niveau_id(l, niveaux_idx),
        } for l in candidat.languesParlees.all()]

    interets = [
        ci.libelleLibre or (ci.typeCentreInteret.nomCentreInteret if ci.typeCentreInteret_id else '')
        for ci in (contenu.interets.all() if contenu else [])
        if ci.libelleLibre or ci.typeCentreInteret_id
    ]
    if not interets:
        interets = [
            ci.libelleLibre or (ci.typeCentreInteret.nomCentreInteret if ci.typeCentreInteret_id else '')
            for ci in candidat.centresInteret.all()
            if ci.libelleLibre or ci.typeCentreInteret_id
        ]

    # `desc` (format CV simplifié) ← `realisation` (= « Vos contributions »).
    # Le modèle `Projet` n'a plus de champ `description` — c'était un orphelin
    # legacy jamais saisi par le formulaire.
    projets = [{
        'id'   : p.id,
        'nom'  : p.titre or '',
        'desc' : p.realisation or '',
        'lien' : p.urlDemo or '',
        'annee': p.dateDebut.year if p.dateDebut else '',
    } for p in (contenu.projets.all() if contenu else [])]
    if not projets:
        projets = [{
            'id'   : p.id,
            'nom'  : p.titre or '',
            'desc' : p.realisation or '',
            'lien' : p.urlDemo or '',
            'annee': p.dateDebut.year if p.dateDebut else '',
        } for p in candidat.projets.all()]

    benevs = [{
        'id'     : b.id,
        'role'   : b.titre or '',
        'orga'   : b.organisation or '',
        'debut'  : _ym(b.dateDebut),
        'fin'    : _ym(b.dateFin),
        'enCours': bool(b.enCours),
    } for b in (contenu.benevolats.all() if contenu else [])]
    if not benevs:
        benevs = [{
            'id'     : b.id,
            'role'   : b.titre or '',
            'orga'   : b.organisation or '',
            'debut'  : _ym(b.dateDebut),
            'fin'    : _ym(b.dateFin),
            'enCours': bool(b.enCours),
        } for b in candidat.benevolats.all()]

    # Application des masques de visibilité (par section, par ID DB)
    _apply_vis(experiences, 'experiences')
    _apply_vis(formations,  'formations')
    _apply_vis(competences, 'competences')
    _apply_vis(langues,     'langues')
    _apply_vis(projets,     'projets')
    _apply_vis(benevs,      'benevs')

    # Centres d'intérêt : strings → indices masqués via tableau parallèle
    interets_masques = list(masques.get('interets') or [])

    return {
        # ── Identifiant interne du CV ──
        'nomCv'    : cv_obj.nomCv or '',
        # ── Identité / contact ──
        'photo'    : _cv_photo_url(cv_obj),
        'prenom'   : prenom,
        'nom'      : nom,
        'titre'    : cv_obj.titre or candidat.titreProfessionnel or '',
        'email'    : email,
        'telephone': telephone,
        'ville'    : ville,
        'pays'     : pays,
        'adresse'  : adresse,
        'age'      : _calculer_age(getattr(candidat, 'dateNaissance', None)),
        'linkedin' : _lien_candidat_par_slug(candidat, 'linkedin'),
        'portfolio': _lien_candidat_par_slug(candidat, 'googlechrome'),
        'permis'   : permis,
        'profil'   : cv_obj.profil or candidat.profilCV or '',

        # ── Collections ──
        'experiences':      experiences,
        'formations':       formations,
        'competences':      competences,
        'langues':          langues,
        'interets':         interets,
        'interetsMasques':  interets_masques,
        'projets':          projets,
        'benevs':           benevs,

        # ── Toggles ──
        'showCertif' : bool(contenu.showCertif)  if contenu else False,
        'showProjets': bool(contenu.showProjets) if contenu else bool(projets),
        'showBenev'  : bool(contenu.showBenev)   if contenu else bool(benevs),
        'showRef'    : bool(contenu.showRef)     if contenu else True,

        # ── Suggestions : items du candidat absents du CV (CV legacy → vide) ──
        'suggestions': {section: [] for section in _SNAPSHOT_SECTIONS},
    }


# ─── Suggestions (diff snapshot ↔ candidat live) ──────────────────────────────

def _signature(section, item):
    """Empreinte texte normalisée d'un item, utilisée comme fallback du diff
    quand l'ID DB ne suffit pas (delete+recreate change les IDs).
    Insensible à la casse / aux espaces. Retourne '' pour items sans contenu
    distinctif (qui n'ont rien à matcher).
    """
    def _norm(value):
        return ' '.join(str(value or '').lower().split()).strip()

    if section == 'experiences':
        return f"{_norm(item.get('entreprise'))}|{_norm((item.get('postes') or [{}])[0].get('titre'))}"
    if section == 'formations':
        return f"{_norm(item.get('diplome'))}|{_norm(item.get('ecole'))}"
    if section == 'competences':
        return _norm(item.get('nom'))
    if section == 'langues':
        return _norm(item.get('nom'))
    if section == 'interets':
        return _norm(item if isinstance(item, str) else item.get('libelle') or item.get('nom'))
    if section == 'projets':
        return _norm(item.get('nom') or item.get('titre'))
    if section == 'benevs':
        return _norm(f"{item.get('role')}|{item.get('orga')}")
    return ''


def _merge_cv_into_profile(candidat, cv_data):
    """Construit le dict « supremum » à persister dans le profil candidat.

    Politique : sync **additif**. Le profil candidat (pool maître) ne perd
    jamais de données quand on sauvegarde un CV. Pour chaque section :
      - on garde les items du CV en premier (ordre + alignement préservés
        pour le calcul de `elementsMasques` qui matche par index),
      - on append les items du profil qui ne sont pas déjà dans le CV
        (matching hybride ID puis empreinte texte).

    Conséquence : à la prochaine création de CV, `build_cv_initial` voit
    bien l'ensemble complet. Les items absents du CV courant deviennent
    automatiquement des suggestions au prochain reload du même CV.
    """
    from .cv_initial import build_rubriques_initial

    profil = build_rubriques_initial(candidat)
    out = dict(cv_data)

    for section in _SNAPSHOT_SECTIONS:
        cv_items   = list(cv_data.get(section) or [])
        prof_items = profil.get(section) or []

        cv_ids = {it.get('id') for it in cv_items
                  if isinstance(it, dict) and it.get('id') is not None}
        cv_sigs = {_signature(section, it) for it in cv_items
                   if _signature(section, it)}

        merged = cv_items
        for prof_it in prof_items:
            prof_id  = prof_it.get('id') if isinstance(prof_it, dict) else None
            prof_sig = _signature(section, prof_it)
            if prof_id is not None and prof_id in cv_ids:
                continue
            if prof_sig and prof_sig in cv_sigs:
                continue
            merged.append(prof_it)
        out[section] = merged

    return out


def _compute_suggestions(candidat, snapshot):
    """Calcule les suggestions à proposer dans l'éditeur.

    Pour chaque section, on identifie les items présents dans le candidat
    (table relationnelle / snapshot Profil) et absents du CV figé. Le matching
    est hybride : ID DB d'abord, puis empreinte texte normalisée si l'ID ne
    correspond pas (cas du delete+recreate par `sync_rubriques_to_db`).

    Retourne `{section: [item, ...]}` au format éditeur (mêmes shapes que les
    listes consommées par les `x-for` Alpine.js du form panel).
    """
    from .cv_initial import build_rubriques_initial

    live = build_rubriques_initial(candidat)
    snap_sections = {s: snapshot.get(s) or [] for s in _SNAPSHOT_SECTIONS}

    out = {}
    for section in _SNAPSHOT_SECTIONS:
        snap_items = snap_sections[section]
        live_items = live.get(section) or []

        snap_ids = {it.get('id') for it in snap_items
                    if isinstance(it, dict) and it.get('id') is not None}
        snap_sigs = {_signature(section, it) for it in snap_items
                     if _signature(section, it)}

        suggestions = []
        for item in live_items:
            item_id  = item.get('id') if isinstance(item, dict) else None
            item_sig = _signature(section, item)
            # Match par ID puis par contenu (hybride)
            if item_id is not None and item_id in snap_ids:
                continue
            if item_sig and item_sig in snap_sigs:
                continue
            suggestions.append(item)
        out[section] = suggestions
    return out


# ─── Génération PDF / PNG / JPG ──────────────────────────────────────────────
# Tout passe par Playwright (Chromium headless) pour un rendu strictement
# identique à l'aperçu navigateur. Voir `candidat/cv_render.py`.


# ─── Génération DOCX ───────────────────────────────────────────────────────────

def _generate_docx(cv):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    ORANGE = RGBColor(0xF7, 0x7F, 0x00)
    GRAY   = RGBColor(0x6B, 0x72, 0x80)

    doc = Document()

    # ── Mise en page A4 ───────────────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2)
    sec.top_margin  = sec.bottom_margin = Cm(1.8)

    # Supprimer les styles de base inutiles
    for style in ('Normal',):
        try:
            doc.styles[style].font.name = 'Arial'
            doc.styles[style].font.size = Pt(10)
        except Exception:
            pass

    def add_heading(title, first=False):
        p = doc.add_paragraph()
        if not first:
            p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(8.5)
        run.font.color.rgb = ORANGE
        # Bordure inférieure orange
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single')
        bot.set(qn('w:sz'), '8')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'F77F00')
        pBdr.append(bot)
        pPr.append(pBdr)
        return p

    def add_item_header(title, badge=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(11)
        if badge:
            tab = p.add_run(f'    {badge}')
            tab.font.size = Pt(8)
            tab.font.color.rgb = ORANGE
        return p

    def add_sub(text, color=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(text)
        r.font.size = Pt(9)
        if color:
            r.font.color.rgb = color
        return p

    # ── En-tête : Nom / Titre ─────────────────────────────────────────────
    # Barre orange en haut
    from docx.oxml import OxmlElement as OE
    from docx.oxml.ns import qn as Qn

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(
        f"{cv.get('prenom', '')} {cv.get('nom', '').upper()}"
    )
    name_run.bold = True
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = ORANGE

    if cv.get('titre'):
        t = doc.add_paragraph()
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = t.add_run(cv['titre'])
        r.italic = True
        r.font.size = Pt(12)
        r.font.color.rgb = GRAY

    # Contacts sur une ligne
    contacts = []
    if cv.get('email'):     contacts.append(f"✉ {cv['email']}")
    if cv.get('telephone'): contacts.append(f"☎ {cv['telephone']}")
    if cv.get('adresse'):   contacts.append(f"📍 {cv['adresse']}")
    if cv.get('age'):       contacts.append(f"📅 {cv['age']} ans")
    if cv.get('linkedin'):  contacts.append(f"💼 {cv['linkedin']}")
    if cv.get('portfolio'): contacts.append(f"🔗 {cv['portfolio']}")
    if contacts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run('  │  '.join(contacts))
        cr.font.size = Pt(9)
        cr.font.color.rgb = GRAY

    # Séparateur
    sep = doc.add_paragraph()
    sep.paragraph_format.space_after = Pt(6)

    # ── Profil ────────────────────────────────────────────────────────────
    if cv.get('profil'):
        add_heading('Profil', first=True)
        p = doc.add_paragraph(cv['profil'])
        p.runs[0].font.size = Pt(10)
        p.paragraph_format.space_after = Pt(2)

    # ── Expériences ───────────────────────────────────────────────────────
    exps = cv.get('experiences_ok', [])
    if exps:
        add_heading('Expériences professionnelles')
        for e in exps:
            postes = e.get('postes', [])
            if len(postes) <= 1:
                add_item_header(e.get('titre', '—'), badge=e.get('periode') or None)
                sub_parts = []
                if e.get('entreprise'): sub_parts.append(e['entreprise'])
                if e.get('lieu'):       sub_parts.append(e['lieu'])
                if sub_parts:
                    add_sub(' · '.join(sub_parts), color=ORANGE)
            else:
                header_parts = list(filter(None, [e.get('entreprise', '—'), e.get('lieu')]))
                add_item_header(' · '.join(header_parts) or '—', badge=e.get('periode') or None)
                for p in postes:
                    if p.get('titre'):
                        pp = doc.add_paragraph()
                        pp.paragraph_format.space_before = Pt(2)
                        pp.paragraph_format.space_after  = Pt(1)
                        pp.paragraph_format.left_indent  = Cm(0.5)
                        r1 = pp.add_run(f"→ {p['titre']}")
                        r1.bold = True
                        r1.font.size = Pt(9)
                        if p.get('periode'):
                            r2 = pp.add_run(f"  {p['periode']}")
                            r2.font.size = Pt(8)
                            r2.font.color.rgb = ORANGE
            mc_list = e.get('missionsClient_ok', []) if e.get('hasMissionsClient') else []
            if mc_list:
                mch = doc.add_paragraph()
                mch.paragraph_format.space_before = Pt(5)
                mch.paragraph_format.space_after = Pt(2)
                mch_r = mch.add_run('Missions clients')
                mch_r.bold = True
                mch_r.font.size = Pt(8.5)
                mch_r.font.color.rgb = ORANGE
                for mc in mc_list:
                    add_item_header(mc.get('client', '—'), badge=mc.get('periode', '') or None)
                    mc_sub = []
                    if mc.get('pays'):  mc_sub.append(mc['pays'])
                    if mc.get('ville'): mc_sub.append(mc['ville'])
                    if mc_sub:
                        add_sub(' · '.join(mc_sub), color=GRAY)
                    if mc.get('desc'):
                        add_sub(mc['desc'])

    # ── Formations (diplômes, attestations, formations libres) ────────────────
    diplomes = cv.get('diplomes_ok', [])
    if diplomes:
        add_heading('Formations')
        for f in diplomes:
            add_item_header(
                f.get('diplome', '—'),
                badge=f.get('periode', '') or None,
            )
            sub_parts = []
            if f.get('ecole'): sub_parts.append(f['ecole'])
            if f.get('lieu'):  sub_parts.append(f['lieu'])
            if sub_parts:
                add_sub(' · '.join(sub_parts), color=ORANGE)
            if f.get('domaine'):
                add_sub(f['domaine'], color=GRAY)
            if f.get('desc'):
                add_sub(f['desc'])

    # ── Certifications (issues des formations + legacy Divers) ────────────────
    certifs_form   = cv.get('certifs_form_ok', [])
    certifs_legacy = cv.get('certifs_ok', []) if cv.get('showCertif') else []
    all_certifs    = certifs_form + certifs_legacy
    if all_certifs:
        add_heading('Certifications')
        for c in all_certifs:
            is_legacy = 'nom' in c and 'organisme' in c and 'typeSortie' not in c
            titre     = c.get('nom', '') if is_legacy else c.get('diplome', '—')
            organisme = c.get('organisme', '') if is_legacy else c.get('ecole', '')
            badge_txt = (str(c.get('annee', '')) if is_legacy else c.get('periode_certif', '')) or None
            add_item_header(titre, badge=badge_txt)
            if organisme:
                add_sub(organisme, color=ORANGE)
            if not is_legacy:
                if c.get('domaine'):
                    add_sub(c['domaine'], color=GRAY)
                if c.get('numero'):
                    add_sub(f"N° {c['numero']}", color=GRAY)
            if c.get('desc'):
                add_sub(c['desc'])

    # ── Compétences ───────────────────────────────────────────────────────
    comps = cv.get('competences_ok', [])
    if comps:
        add_heading('Compétences')
        for c in comps:
            n = int(c.get('niveau_nb', 0))
            dots = '●' * n + '○' * (5 - n)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{c['nom']}   ")
            r1.font.size = Pt(9.5)
            r2 = p.add_run(dots)
            r2.font.size = Pt(10)
            r2.font.color.rgb = ORANGE

    # ── Langues ───────────────────────────────────────────────────────────
    langs = cv.get('langues_ok', [])
    if langs:
        add_heading('Langues')
        for l in langs:
            n = int(l.get('niveau_nb', 0))
            dots = '●' * n + '○' * (5 - n)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{l['nom']}   ")
            r1.font.size = Pt(9.5)
            r2 = p.add_run(dots)
            r2.font.size = Pt(10)
            r2.font.color.rgb = ORANGE
            label = l.get('niveau_label', '')
            if label:
                r3 = p.add_run(f'  ({label})')
                r3.font.size = Pt(8.5)
                r3.font.color.rgb = GRAY

    # ── Centres d'intérêt ─────────────────────────────────────────────────
    interets = cv.get('interets_ok', [])
    if interets:
        add_heading("Centres d'intérêt")
        p = doc.add_paragraph(' · '.join(interets))
        p.runs[0].font.size = Pt(9.5)

    # ── Certifications legacy déjà traitées dans le bloc Certifications ci-dessus

    # ── Projets ───────────────────────────────────────────────────────────
    if cv.get('showProjets'):
        projets = cv.get('projets_ok', [])
        if projets:
            add_heading('Projets personnels')
            for p in projets:
                add_item_header(p.get('nom', ''), badge=str(p['annee']) if p.get('annee') else None)
                if p.get('desc'):  add_sub(p['desc'])
                if p.get('lien'):
                    lr = doc.add_paragraph()
                    lr.paragraph_format.space_after = Pt(2)
                    r = lr.add_run(f"🔗 {p['lien']}")
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = ORANGE

    # ── Bénévolat ─────────────────────────────────────────────────────────
    if cv.get('showBenev'):
        benevs = cv.get('benevs_ok', [])
        if benevs:
            add_heading('Bénévolat / Associations')
            for b in benevs:
                add_item_header(b.get('role', ''), badge=b.get('periode') or None)
                if b.get('orga'): add_sub(b['orga'], color=GRAY)

    # ── Références ────────────────────────────────────────────────────────
    if cv.get('showRef'):
        rp = doc.add_paragraph()
        rp.paragraph_format.space_before = Pt(14)
        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = rp.add_run('Références disponibles sur demande')
        r.italic = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = GRAY

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ─── Vues CV ──────────────────────────────────────────────────────────────────

def modeles_cv(request):
    from django.core.paginator import Paginator

    cache_key = 'modeles_cv_data'
    cached = django_cache.get(cache_key)
    if cached is None:
        modeles_qs = ModeleCV.objects.filter(actif=True).order_by('ordre', 'nom')
        cached = {
            'modeles_list': list(modeles_qs),
            'categories':   ['Tous'] + [label for _, label in ModeleCV.CATEGORIES],
            'secteurs':     [label for _, label in ModeleCV.SECTEURS],
            'total':        modeles_qs.count(),
            'gratuits':     modeles_qs.filter(premium=False).count(),
            'premium':      modeles_qs.filter(premium=True).count(),
        }
        django_cache.set(cache_key, cached, 1800)

    categories = cached['categories']
    secteurs   = cached['secteurs']
    total      = cached['total']
    gratuits   = cached['gratuits']
    premium    = cached['premium']

    PER_PAGE = 20
    page_number = request.GET.get('page', 1)
    paginator = Paginator(cached['modeles_list'], PER_PAGE)
    page_obj  = paginator.get_page(page_number)

    cur   = page_obj.number
    total_pages = paginator.num_pages
    start = max(1, cur - 2)
    end   = min(total_pages, cur + 2)
    if end - start < 4:
        if start == 1:
            end = min(total_pages, start + 4)
        else:
            start = max(1, end - 4)
    page_range = range(start, end + 1)

    params = request.GET.copy()
    params.pop('page', None)
    qs = params.urlencode()

    return render(request, 'candidat/modeles_cv.html', {
        'modeles':    page_obj,
        'categories': categories,
        'secteurs':   secteurs,
        'total':      total,
        'gratuits':   gratuits,
        'premium':    premium,
        'page_obj':   page_obj,
        'page_range': page_range,
        'qs':         qs,
    })


def apercu_cv(request, template_id):
    """Prévisualisation du modèle de CV (sans formulaire, données de démo). Accessible sans connexion."""
    modele = ModeleCV.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return redirect('candidat:modeles_cv')
    modeles = ModeleCV.objects.filter(actif=True).order_by('ordre', 'nom')
    from .niveau_resolver import niveaux_for_editor
    return render(request, f"candidat/cv/modeles/{modele.fichier}.html", {
        'modele':      modele,
        'modeles':     modeles,
        'apercu':      True,
        'niveaux_ref': niveaux_for_editor(),
    })


@candidat_required
def creer_cv(request, template_id):
    modele = ModeleCV.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        messages.error(request, 'Ce modèle est introuvable ou n\'est plus disponible.')
        return redirect('candidat:modeles_cv')
    if modele.premium:
        messages.error(request, 'Ce modèle est réservé aux membres Premium.')
        return redirect('candidat:modeles_cv')

    from .cv_initial import build_cv_initial
    from .niveau_resolver import niveaux_for_editor
    cv_initial = build_cv_initial(request.candidat)

    modeles = ModeleCV.objects.filter(actif=True).order_by('ordre', 'nom')
    return render(request, f"candidat/cv/modeles/{modele.fichier}.html", {
        'modele':       modele,
        'modeles':      modeles,
        'cv_initial':   cv_initial,
        'candidat_id':  request.candidat.id,
        'cv_id':        None,
        'niveaux_ref':  niveaux_for_editor(),
    })


# ─── Fragment d'aperçu (switch de modèle sans reload) ────────────────────────

class _CvFragmentExtractor:
    """Extrait `<style>` et le contenu de `<div id="cv-sheet">` d'un HTML rendu.

    Utilise html.parser (stdlib) — pas de dépendance externe. Robuste aux divs
    imbriqués grâce à un compteur de profondeur.
    """

    def __init__(self):
        from html.parser import HTMLParser

        outer = self

        class _Parser(HTMLParser):
            def __init__(self):
                super().__init__(convert_charrefs=False)

            def handle_starttag(self, tag, attrs):
                if tag.lower() == 'style':
                    outer._in_style  = True
                    outer._style_buf = []
                    return
                attrs_d = dict(attrs)
                if tag.lower() == 'div' and attrs_d.get('id') == 'cv-sheet':
                    outer._depth = 1
                    outer._html_parts.append(self.get_starttag_text())
                    return
                if outer._depth > 0:
                    if tag.lower() == 'div':
                        outer._depth += 1
                    outer._html_parts.append(self.get_starttag_text())

            def handle_endtag(self, tag):
                if tag.lower() == 'style' and outer._in_style:
                    outer.styles.append(''.join(outer._style_buf))
                    outer._in_style = False
                    outer._style_buf = []
                    return
                if outer._depth > 0:
                    if tag.lower() == 'div':
                        outer._depth -= 1
                        outer._html_parts.append('</div>')
                        if outer._depth == 0:
                            return
                    else:
                        outer._html_parts.append(f'</{tag}>')

            def handle_startendtag(self, tag, attrs):
                if outer._depth > 0:
                    outer._html_parts.append(self.get_starttag_text())

            def handle_data(self, data):
                if outer._in_style:
                    outer._style_buf.append(data)
                elif outer._depth > 0:
                    outer._html_parts.append(data)

            def handle_entityref(self, name):
                if outer._depth > 0:
                    outer._html_parts.append(f'&{name};')
                elif outer._in_style:
                    outer._style_buf.append(f'&{name};')

            def handle_charref(self, name):
                if outer._depth > 0:
                    outer._html_parts.append(f'&#{name};')
                elif outer._in_style:
                    outer._style_buf.append(f'&#{name};')

            def handle_comment(self, data):
                if outer._depth > 0:
                    outer._html_parts.append(f'<!--{data}-->')

        self.styles      = []
        self._style_buf  = []
        self._in_style   = False
        self._html_parts = []
        self._depth      = 0
        self._parser     = _Parser()

    def feed(self, html):
        self._parser.feed(html)
        return self

    @property
    def html(self):
        return ''.join(self._html_parts)

    @property
    def css(self):
        return '\n'.join(s for s in self.styles if s.strip())


@candidat_required
def apercu_modele_fragment(request, template_id):
    """Retourne CSS + HTML de l'aperçu d'un modèle, pour switch sans reload.

    Le frontend (`changerModele()`) consomme ce JSON pour remplacer uniquement
    le `<div id="cv-sheet">` et le bloc CSS du modèle, sans recharger la page
    ni perdre l'état Alpine `cv` (saisies utilisateur).
    """
    modele = ModeleCV.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return JsonResponse({'ok': False, 'message': 'Modèle introuvable.'}, status=404)
    if modele.premium:
        return JsonResponse({'ok': False, 'message': 'Modèle réservé aux abonnés Premium.'}, status=403)

    from .cv_initial import build_cv_initial
    cv_initial = build_cv_initial(request.candidat)

    html = render_to_string(
        f"candidat/cv/modeles/{modele.fichier}.html",
        {
            'modele':      modele,
            'modeles':     ModeleCV.objects.filter(actif=True).order_by('ordre', 'nom'),
            'cv_initial':  cv_initial,
            'candidat_id': request.candidat.id,
            'cv_id':       None,
        },
        request=request,
    )

    extractor = _CvFragmentExtractor().feed(html)

    return JsonResponse({
        'ok':         True,
        'modele_id':  modele.id,
        'modele_nom': modele.nom,
        'css':        extractor.css,
        'html':       extractor.html,
    })


@candidat_required
def modifier_cv(request, cv_id):
    """Ouvre l'éditeur sur un CV déjà sauvegardé pour modification.

    Reconstruit l'état initial 100 % depuis l'ORM via `_cv_to_dict(cv)` :
    identité depuis le snapshot `CVContenu.informationPersonnelle`, rubriques
    depuis les M2M de `CVContenu` (avec fallback sur les rubriques live du
    candidat si une M2M est vide). Le `cv_id` est injecté dans le builder
    Alpine → prochain save = update du même CV.
    """
    cv = (
        CV.objects.filter(pk=cv_id, candidat=request.candidat, archive=False)
        .select_related('modele', 'contenu', 'contenu__informationPersonnelle')
        .first()
    )
    if not cv or not cv.modele or not cv.modele.actif:
        messages.error(request, 'CV introuvable ou modèle indisponible.')
        return redirect('candidat:profil')

    cv_initial = _cv_to_dict(cv)

    modeles = ModeleCV.objects.filter(actif=True).order_by('ordre', 'nom')
    from .niveau_resolver import niveaux_for_editor
    return render(request, f"candidat/cv/modeles/{cv.modele.fichier}.html", {
        'modele':       cv.modele,
        'modeles':      modeles,
        'cv_initial':   cv_initial,
        'candidat_id':  request.candidat.id,
        'cv_id':        cv.pk,
        'niveaux_ref':  niveaux_for_editor(),
    })


@candidat_required
@require_POST
def archiver_cv(request, cv_id):
    """Archive (soft-delete) un CV appartenant au candidat connecté."""
    cv = CV.objects.filter(pk=cv_id, candidat=request.candidat).first()
    if not cv:
        return JsonResponse({'ok': False, 'message': 'CV introuvable.'}, status=404)
    cv.archiver()
    return JsonResponse({'ok': True, 'message': 'CV archivé.'})


@candidat_required
@require_POST
def regenerer_cv_artefacts(request, cv_id):
    """Régénère les artefacts (PDF + images) d'un CV existant.

    Comportement attendu : agir EXACTEMENT comme le bouton « Sauvegarder »
    côté éditeur pour ce qui concerne les artefacts — c.-à-d. :
      • on reconstruit `cv_data` depuis l'ORM via `_cv_to_dict(cv)` (état
        à jour du CV après modification éventuelle du modèle ou des rubriques) ;
      • on appelle `cv.generer_artefacts(...)` qui ÉCRASE :
          - `cv.cvPdf` (FileField.save = remplace le fichier)
          - toutes les `PhotoCV` existantes (delete + recreate)
      • on renvoie les nouvelles URLs PDF + 1ʳᵉ image pour rafraîchir l'UI.

    Aucun nouvel enregistrement CV n'est créé — c'est bien une mise à jour
    in-place dans l'historique.
    """
    cv = (
        CV.objects.filter(pk=cv_id, candidat=request.candidat, archive=False)
        .select_related('modele', 'contenu', 'contenu__informationPersonnelle')
        .first()
    )
    if not cv:
        return JsonResponse({'ok': False, 'message': 'CV introuvable.'}, status=404)
    if not cv.modele or not cv.modele.actif:
        return JsonResponse(
            {'ok': False, 'message': 'Modèle indisponible — impossible de régénérer.'},
            status=400,
        )

    try:
        cv_data = _cv_to_dict(cv)
        ok = cv.generer_artefacts(request, cv_data=cv_data)
    except Exception as exc:
        logger.exception('Échec régénération artefacts CV %s', cv_id)
        return JsonResponse(
            {'ok': False, 'message': f"Erreur lors de la régénération : {exc}"},
            status=500,
        )

    if not ok:
        return JsonResponse(
            {'ok': False, 'message': "La régénération a échoué (modèle ou contenu incomplet)."},
            status=500,
        )

    # Renvoie les URLs fraîchement écrasées pour rafraîchir la carte côté UI
    photo = cv.photos.order_by('numeroPage').first()
    return JsonResponse({
        'ok':            True,
        'message':       'CV mis à jour — PDF et images régénérés.',
        'pdf_url':       cv.cvPdf.url   if cv.cvPdf   else '',
        'image_url':     photo.image.url if (photo and photo.image) else '',
        'thumbnail_url': photo.image.url if (photo and photo.image) else '',
        'nb_pages':      cv.photos.count(),
    })


@candidat_required
def api_lister_cvs(request):
    """JSON : liste des CV non archivés du candidat connecté.

    Sert d'API au store Alpine `cvList` côté profil (onglet « Mes CV ») qui
    rafraîchit la liste à chaque entrée d'onglet et après chaque action.

    `pdf_pending` est `True` quand les artefacts (PDF + images) ne sont pas
    encore matérialisés ET que la modification est récente (< 2 min) — donc
    probablement en cours de génération asynchrone. Le frontend re-poll
    jusqu'à apparition du PDF, sans intervention de l'utilisateur.
    """
    from django.utils import timezone
    from datetime import timedelta

    cvs = (
        request.candidat.cvs.filter(archive=False)
        .select_related('modele')
        .prefetch_related('photos')
        .order_by('-dateModification')
    )

    pending_threshold = timezone.now() - timedelta(seconds=120)

    def _serialize(cv):
        # Utilise la liste préchargée (prefetch_related) — évite N requêtes count/first
        photos     = list(cv.photos.all())
        miniature  = photos[0] if photos else None
        image_url  = miniature.image.url if miniature and miniature.image else ''
        nb_pages   = len(photos)
        has_pdf    = bool(cv.cvPdf)
        # Artefacts manquants + modif récente → génération en cours côté worker.
        # Au-delà de 120s sans artefact, on considère que la génération a échoué
        # (Playwright KO, etc.) — on ne maintient pas un poll infini.
        pdf_pending = (not has_pdf or not image_url) and cv.dateModification > pending_threshold
        return {
            'id':                cv.pk,
            'nomCv':             cv.nomCv or '',
            'titre':             cv.titre or 'CV sans titre',
            'modele_nom':        cv.modele.nom if cv.modele else '',
            'modele_indispo':    cv.modele is None,
            'date_modification': cv.dateModification.strftime('%d/%m/%Y à %H:%M'),
            'modifier_url':      reverse('candidat:modifier_cv', args=[cv.pk]) if cv.modele else '',
            'pdf_url':           cv.cvPdf.url if has_pdf else '',
            'image_url':         image_url,    # 1ʳᵉ page — sert uniquement de miniature
            'thumbnail_url':     image_url,
            'nb_pages':          nb_pages,
            # Endpoint qui sert PNG (1 page) ou ZIP (plusieurs pages) selon nb_pages
            'images_zip_url':    reverse('candidat:api_images_cv', args=[cv.pk]) if image_url else '',
            'pdf_pending':       pdf_pending,
            'est_importe':       cv.estImporte,
        }

    return JsonResponse({'ok': True, 'cvs': [_serialize(c) for c in cvs]})


# ── Téléchargement images multi-pages ────────────────────────────────────────

@candidat_required
def api_images_cv(request, cv_id):
    """Télécharge les pages d'un CV en image.

    • 1 page  → fichier PNG direct.
    • N pages → archive ZIP contenant N fichiers PNG (une image par page).

    Le nom du fichier est déduit de `cv.nomCv` ou `cv.titre`.
    """
    cv = CV.objects.filter(pk=cv_id, candidat=request.candidat, archive=False).first()
    if not cv:
        return HttpResponse('CV introuvable.', status=404)

    photos = list(cv.photos.order_by('numeroPage'))
    if not photos:
        return HttpResponse('Aucune image disponible.', status=404)

    nom_base = cv.nomCv or cv.titre or f'CV_{cv.pk}'
    # Nettoyage : retire les caractères interdits dans les noms de fichiers
    nom_base = ''.join(c for c in nom_base if c not in r'\/:*?"<>|')

    if len(photos) == 1:
        photo = photos[0]
        response = HttpResponse(photo.image.read(), content_type='image/png')
        response['Content-Disposition'] = f'attachment; filename="{nom_base}.png"'
        return response

    # Plusieurs pages → ZIP
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for photo in photos:
            zf.writestr(f'{nom_base}_page{photo.numeroPage}.png', photo.image.read())
    buf.seek(0)
    response = HttpResponse(buf.getvalue(), content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{nom_base}_images.zip"'
    return response


# ── Helpers extraction ────────────────────────────────────────────────────────

def _parse_date_extraction(s):
    """Convertit 'YYYY-MM' ou 'YYYY' en date Python (jour=1). Retourne None si invalide."""
    if not s:
        return None
    from datetime import date as _date
    try:
        parts = str(s).strip().split('-')
        annee = int(parts[0])
        mois  = max(1, min(12, int(parts[1]))) if len(parts) >= 2 else 1
        return _date(annee, mois, 1)
    except (ValueError, TypeError, IndexError):
        return None


# ── Vues : import / extraction ────────────────────────────────────────────────

@candidat_required
@require_POST
def api_importer_cv(request):
    """Importe un PDF depuis l'appareil du candidat et le stocke comme CV sans modèle.

    • Valide le fichier (magic bytes %PDF, taille ≤ 10 Mo).
    • Vérifie l'unicité du nom par candidat.
    • Génère la miniature (toutes pages) via PyMuPDF (_pdf_to_pages).
    • Crée un CV(estImporte=True, modele=None, contenu=None).
    """
    fichier = request.FILES.get('fichier')
    nom_cv  = request.POST.get('nom_cv', '').strip()

    if not fichier:
        return JsonResponse({'ok': False, 'message': 'Aucun fichier reçu.'}, status=400)
    if not nom_cv:
        return JsonResponse({'ok': False, 'message': 'Le nom du CV est obligatoire.'}, status=400)
    if fichier.size > 10 * 1024 * 1024:
        return JsonResponse({'ok': False, 'message': 'Fichier trop volumineux (max 10 Mo).'}, status=400)

    pdf_bytes = fichier.read()
    if not pdf_bytes.startswith(b'%PDF'):
        return JsonResponse({'ok': False, 'message': 'Le fichier doit être un PDF valide.'}, status=400)

    if CV.objects.filter(candidat=request.candidat, nomCv=nom_cv, archive=False).exists():
        return JsonResponse(
            {'ok': False, 'message': f'Vous avez déjà un CV nommé « {nom_cv} ».'},
            status=400,
        )

    try:
        with transaction.atomic():
            cv = CV.objects.create(
                candidat   = request.candidat,
                modele     = None,
                contenu    = None,
                estImporte = True,
                nomCv      = nom_cv,
            )
            cv.cvPdf.save(
                f'import_{request.candidat.pk}_{cv.pk}.pdf',
                ContentFile(pdf_bytes),
                save=True,
            )
            # Miniature — best-effort : si PyMuPDF ou le PDF pose problème, on continue
            try:
                from .cv_render import _pdf_to_pages
                pages = _pdf_to_pages(pdf_bytes, fmt='png')
                for i, page_bytes in enumerate(pages, start=1):
                    photo = PhotoCV(cv=cv, numeroPage=i)
                    photo.image.save(
                        f'import_{request.candidat.pk}_{cv.pk}_p{i}.png',
                        ContentFile(page_bytes),
                        save=True,
                    )
            except Exception:
                logger.warning('Miniature non générée pour CV importé #%s', cv.pk)

    except Exception as exc:
        logger.exception("Échec import CV candidat #%s", request.candidat.pk)
        return JsonResponse({'ok': False, 'message': f"Erreur lors de l'import : {exc}"}, status=500)

    return JsonResponse({'ok': True, 'message': 'CV importé avec succès.', 'cv_id': cv.pk})


@candidat_required
@require_POST
def api_extraire_cv(request, cv_id):
    """Extrait les données d'un CV via Ollama en streaming SSE.

    Envoie des événements Server-Sent Events au format :
      data: {"tokens": N}          — progression (tous les 10 tokens)
      data: {"ok": true,  "donnees": {...}}  — résultat final
      data: {"ok": false, "message": "..."}  — erreur
    """
    from django.http import StreamingHttpResponse

    cv = CV.objects.filter(pk=cv_id, candidat=request.candidat, estImporte=True, archive=False).first()
    if not cv:
        return JsonResponse({'ok': False, 'message': 'CV introuvable.'}, status=404)
    if not cv.cvPdf:
        return JsonResponse({'ok': False, 'message': 'Fichier PDF introuvable.'}, status=400)

    # ── Lecture du PDF ───────────────────────────────────────────────────────
    try:
        import fitz
        doc   = fitz.open(stream=cv.cvPdf.read(), filetype='pdf')
        texte = '\n'.join(page.get_text() for page in doc)
        doc.close()
    except Exception as exc:
        logger.exception("Lecture PDF échouée pour CV importé #%s", cv_id)
        return JsonResponse({'ok': False, 'message': f'Impossible de lire le PDF : {exc}'}, status=500)

    if not texte.strip():
        return JsonResponse(
            {'ok': False, 'message': 'Ce PDF ne contient pas de texte extractible (PDF scanné ?).'},
            status=400,
        )

    # ── Import Ollama ────────────────────────────────────────────────────────
    try:
        import ollama as _ollama
    except ImportError:
        return JsonResponse(
            {'ok': False, 'message': 'Ollama n\'est pas installé. Lancez : pip install ollama'},
            status=503,
        )

    prompt = (
        "Extrais les informations du CV suivant. "
        "Retourne un objet JSON avec exactement ces clés : "
        "titre (string), profil (string), "
        "formations (liste de {diplome,ecole,ville,pays,debut,fin,en_cours,description}), "
        "experiences (liste de {entreprise,ville,pays,debut,fin,en_cours,poste,description}), "
        "competences (liste de {nom,niveau} niveau=1-5), "
        "langues (liste de {nom,niveau} niveau=A1/A2/B1/B2/C1/C2), "
        "interets (liste de {nom}). "
        "Dates au format YYYY-MM.\n\n"
        f"CV :\n{texte[:3500]}"
    )

    def _stream():
        brut       = ''
        nb_tokens  = 0
        try:
            client = _ollama.Client()
            for chunk in client.chat(
                model    = 'mistral',
                messages = [{'role': 'user', 'content': prompt}],
                format   = 'json',
                stream   = True,
                options  = {'temperature': 0, 'num_predict': 800, 'num_ctx': 4096},
            ):
                token = chunk['message']['content']
                brut += token
                nb_tokens += 1
                if nb_tokens % 10 == 0:
                    yield f"data: {json.dumps({'tokens': nb_tokens})}\n\n"
        except Exception as exc:
            logger.exception("Ollama streaming échoué pour CV #%s", cv_id)
            yield f"data: {json.dumps({'ok': False, 'message': f'Ollama inaccessible : {exc}'})}\n\n"
            return

        # Nettoyage éventuel des balises markdown
        if '```' in brut:
            debut = brut.find('{')
            fin   = brut.rfind('}') + 1
            if debut >= 0 and fin > debut:
                brut = brut[debut:fin]

        try:
            donnees = json.loads(brut)
            yield f"data: {json.dumps({'ok': True, 'donnees': donnees})}\n\n"
        except json.JSONDecodeError:
            logger.warning("JSON invalide reçu d'Ollama pour CV #%s : %.200s", cv_id, brut)
            msg = "L'extraction a produit un résultat invalide. Réessayez."
            yield f"data: {json.dumps({'ok': False, 'message': msg})}\n\n"

    resp = StreamingHttpResponse(_stream(), content_type='text/event-stream; charset=utf-8')
    resp['Cache-Control']    = 'no-cache'
    resp['X-Accel-Buffering'] = 'no'
    return resp


@candidat_required
@require_POST
def api_sauvegarder_extraction(request):
    """Enregistre les données extraites par Ollama dans le profil du candidat.

    Les rubriques sont ajoutées (append), jamais écrasées — le candidat
    conserve ses données existantes et peut supprimer les doublons depuis
    son profil.
    """
    try:
        donnees = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'ok': False, 'message': 'Corps de requête JSON invalide.'}, status=400)

    candidat = request.candidat

    try:
        with transaction.atomic():
            # ── Titre & profil ───────────────────────────────────────────────
            champs = {}
            if donnees.get('titre'):
                champs['titreProfessionnel'] = str(donnees['titre'])[:255]
            if donnees.get('profil'):
                champs['profilCV'] = str(donnees['profil'])
            if champs:
                type(candidat).objects.filter(pk=candidat.pk).update(**champs)

            # ── Formations ───────────────────────────────────────────────────
            for f in (donnees.get('formations') or []):
                diplome = str(f.get('diplome') or '').strip()[:200]
                ecole   = str(f.get('ecole')   or '').strip()[:200]
                if not diplome and not ecole:
                    continue
                Formation.objects.create(
                    candidat    = candidat,
                    diplomeLibre= diplome,
                    ecoleLibre  = ecole,
                    ville       = str(f.get('ville') or '')[:100],
                    paysLibre   = str(f.get('pays')  or '')[:100],
                    dateDebut   = _parse_date_extraction(f.get('debut')),
                    dateFin     = _parse_date_extraction(f.get('fin')),
                    enCours     = bool(f.get('en_cours', False)),
                    description = str(f.get('description') or ''),
                )

            # ── Expériences ──────────────────────────────────────────────────
            for e in (donnees.get('experiences') or []):
                entreprise = str(e.get('entreprise') or '').strip()[:200]
                poste      = str(e.get('poste')      or '').strip()[:200]
                if not entreprise and not poste:
                    continue
                exp = ExperienceProfessionnelle.objects.create(
                    candidat       = candidat,
                    entrepriseLibre= entreprise,
                    ville          = str(e.get('ville') or '')[:150],
                    paysLibre      = str(e.get('pays')  or '')[:100],
                    dateDebut      = _parse_date_extraction(e.get('debut')),
                    dateFin        = _parse_date_extraction(e.get('fin')),
                    enCours        = bool(e.get('en_cours', False)),
                    description    = str(e.get('description') or ''),
                )
                if poste:
                    PosteOccupe.objects.create(
                        experience = exp,
                        titreLibre = poste,
                        dateDebut  = _parse_date_extraction(e.get('debut')),
                        dateFin    = _parse_date_extraction(e.get('fin')),
                        enCours    = bool(e.get('en_cours', False)),
                    )

            # ── Compétences ──────────────────────────────────────────────────
            for c in (donnees.get('competences') or []):
                nom = str(c.get('nom') or '').strip()[:200]
                if not nom:
                    continue
                try:
                    etoiles = max(1, min(5, int(c.get('niveau', 3))))
                except (TypeError, ValueError):
                    etoiles = 3
                Competence.objects.create(
                    candidat      = candidat,
                    nomLibre      = nom,
                    valeurEtoiles = etoiles,
                )

            # ── Langues ──────────────────────────────────────────────────────
            for l in (donnees.get('langues') or []):
                nom = str(l.get('nom') or '').strip()[:100]
                if not nom:
                    continue
                CandidatLangue.objects.create(
                    candidat   = candidat,
                    nomLibre   = nom,
                    niveauCode = str(l.get('niveau') or '')[:10],
                )

            # ── Centres d'intérêt ────────────────────────────────────────────
            for i in (donnees.get('interets') or []):
                nom = str(i.get('nom') or '').strip()[:150]
                if not nom:
                    continue
                CentreInteret.objects.create(
                    candidat     = candidat,
                    libelleLibre = nom,
                )

    except Exception as exc:
        logger.exception("Échec sauvegarde extraction CV candidat #%s", candidat.pk)
        return JsonResponse({'ok': False, 'message': f'Erreur lors de la sauvegarde : {exc}'}, status=500)

    return JsonResponse({'ok': True, 'message': 'Profil mis à jour avec les données extraites.'})


@candidat_required
@require_POST
def telecharger_cv(request, template_id, fmt):
    """Génère et télécharge le CV en PDF, PNG, JPG ou DOCX.

    PDF / PNG / JPG : Playwright (Chromium headless) — rendu strictement
    identique à l'aperçu navigateur. Voir `cv_render.py`.
    DOCX : python-docx (best-effort, le format ne permet pas un WYSIWYG strict).
    """
    from . import cv_render

    if fmt not in ('pdf', 'png', 'jpg', 'docx'):
        return HttpResponse('Format non supporté', status=400)

    try:
        cv_data = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return HttpResponse('Données invalides', status=400)

    modele = ModeleCV.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return HttpResponse('Modèle introuvable', status=404)

    prenom   = cv_data.get('prenom', 'Mon').strip().replace(' ', '-') or 'Mon'
    nom      = cv_data.get('nom',    'CV').strip().replace(' ', '-')  or 'CV'
    filename = f'CV-{prenom}-{nom}'

    cv_data = _preprocess_cv(cv_data)

    try:
        if fmt == 'pdf':
            data         = cv_render.render_pdf(modele, cv_data, request)
            content_type = 'application/pdf'
            dl_filename  = f'{filename}.pdf'

        elif fmt in ('png', 'jpg'):
            # CV multi-pages → on rend chaque page en image séparée. Si une seule
            # page : fichier unique. Sinon : ZIP avec page-1, page-2, etc.
            _pdf, pages = cv_render.render_pages(modele, cv_data, request, fmt=fmt)
            if len(pages) <= 1:
                data         = pages[0] if pages else b''
                content_type = 'image/png' if fmt == 'png' else 'image/jpeg'
                dl_filename  = f'{filename}.{fmt}'
            else:
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for i, page_bytes in enumerate(pages, start=1):
                        zf.writestr(f'{filename}-page-{i}.{fmt}', page_bytes)
                data         = buf.getvalue()
                content_type = 'application/zip'
                dl_filename  = f'{filename}.zip'

        else:  # docx
            data         = _generate_docx(cv_data)
            content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            dl_filename  = f'{filename}.docx'

    except ImportError as exc:
        return HttpResponse(
            f'Bibliothèque manquante : {exc}\n'
            f'Installez-la : pip install playwright pymupdf pillow python-docx '
            f'&& playwright install chromium',
            status=500, content_type='text/plain; charset=utf-8',
        )
    except Exception as exc:
        logger.exception('Échec génération CV format=%s template=%s', fmt, template_id)
        return HttpResponse(
            f'Erreur de génération : {exc}',
            status=500, content_type='text/plain; charset=utf-8',
        )

    resp = HttpResponse(data, content_type=content_type)
    resp['Content-Disposition'] = f'attachment; filename="{dl_filename}"'
    resp['Content-Length']      = len(data)
    return resp


# ─── Sauvegarde du CV (relationnel pur) ───────────────────────────────────────

def _txt(value):
    return ('' if value is None else str(value)).strip()


def _upsert_information_personnelle(candidat, cv_data):
    """Met à jour (ou crée) l'InformationPersonnelle du candidat depuis le payload.

    Une seule InformationPersonnelle par candidat (OneToOne). Tous les CVs du
    candidat pointent vers cette unique instance via `CVContenu.informationPersonnelle`.
    """
    info = getattr(candidat, 'informationPersonnelle', None)
    if info is None:
        info = InformationPersonnelle.objects.create(
            nom    = _txt(cv_data.get('nom'))    or candidat.nom    or '',
            prenom = _txt(cv_data.get('prenom')) or candidat.prenom or '',
            email  = candidat.email,
        )
        candidat.informationPersonnelle = info
        candidat.save(update_fields=['informationPersonnelle'])
    else:
        info.nom    = _txt(cv_data.get('nom'))    or info.nom
        info.prenom = _txt(cv_data.get('prenom')) or info.prenom

    info.telephone = _txt(cv_data.get('telephone'))
    info.adresse   = _txt(cv_data.get('adresse'))
    info.ville     = _txt(cv_data.get('ville'))
    info.pays      = _txt(cv_data.get('pays'))
    info.permis    = _txt(cv_data.get('permis'))
    info.save()
    return info


def _apply_cv_photo(cv_obj, candidat, photo_data):
    """Applique la photo reçue de l'éditeur au CV (+ au candidat si vide).

    Règles :
      • Si `photo_data` est une data URL base64 → décode et enregistre comme
        photo PROPRE du CV. Si le candidat n'a pas encore de photo de profil
        générale, on en profite pour la lui assigner (one-shot).
      • Si `photo_data` est une URL existante (même photo qu'avant) → no-op.
      • Si `photo_data` est vide / null → on efface la photo du CV mais on ne
        touche pas à la photo de profil du candidat.
    """
    import base64
    from django.core.files.base import ContentFile

    photo_data = (photo_data or '').strip()

    if not photo_data:
        if cv_obj.photoProfil:
            try:
                cv_obj.photoProfil.delete(save=False)
            except Exception:
                pass
        return

    if not photo_data.startswith('data:'):
        return

    try:
        header, b64 = photo_data.split(',', 1)
        mime = header.split(';')[0].split(':', 1)[1] if ':' in header else 'image/png'
        ext = mime.split('/')[-1].split('+')[0] or 'png'
        raw = base64.b64decode(b64)
    except Exception:
        return

    if cv_obj.photoProfil:
        try:
            cv_obj.photoProfil.delete(save=False)
        except Exception:
            pass

    filename = f'cv_{candidat.pk}_photo.{ext}'
    cv_obj.photoProfil.save(filename, ContentFile(raw), save=False)

    candidat_photo = getattr(candidat, 'photoProfil', None)
    if not (candidat_photo and getattr(candidat_photo, 'name', '')):
        candidat.photoProfil.save(filename, ContentFile(raw), save=True)


def _sync_candidat_identity(candidat, cv_data):
    """Met à jour les champs identité/contact directs sur Candidat depuis le payload.

    Champs mis à jour : prenom, nom, telephone, adresse, titreProfessionnel,
    profilCV. L'email N'est PAS modifié (clé d'auth).
    `profilCV` (résumé du CV) est volontairement distinct de `biographie`
    (texte du portfolio) — les deux ne se chevauchent pas.
    Liens sociaux (linkedin / portfolio) : upsert dans la table `LienCandidat`
    via le référentiel `ReseauSocial` — voir `_upsert_lien_candidat`.
    """
    fields_changed = []
    mapping = {
        'prenom'             : _txt(cv_data.get('prenom')),
        'nom'                : _txt(cv_data.get('nom')),
        'telephone'          : _txt(cv_data.get('telephone')),
        'adresse'            : _txt(cv_data.get('adresse')),
        'titreProfessionnel' : _txt(cv_data.get('titre')),
        'profilCV'           : _txt(cv_data.get('profil')),
    }
    for field, value in mapping.items():
        if not value and field in ('prenom', 'nom'):
            # Ne jamais effacer le nom/prénom existants avec une chaîne vide.
            continue
        if getattr(candidat, field) != value:
            setattr(candidat, field, value)
            fields_changed.append(field)
    if fields_changed:
        candidat.save(update_fields=fields_changed)

    _upsert_lien_candidat(candidat, 'linkedin',     _txt(cv_data.get('linkedin')))
    _upsert_lien_candidat(candidat, 'googlechrome', _txt(cv_data.get('portfolio')))


def _lien_candidat_par_slug(candidat, slug):
    """URL du lien social du candidat pour ce slug réseau, ou '' si absent."""
    lien = candidat.liensSociaux.filter(reseau__slug=slug).first()
    return lien.url if lien else ''


def _upsert_lien_candidat(candidat, slug, url):
    """Crée / met à jour / supprime le LienCandidat pour ce réseau.

    URL vide → suppression de la ligne. URL non vide → upsert. Réseau
    inexistant ou inactif → no-op (pas d'erreur, garde-fou).
    """
    from .models import LienCandidat
    from referentiel.models import ReseauSocial

    reseau = ReseauSocial.objects.filter(slug=slug, actif=True).first()
    if not reseau:
        return
    if url:
        LienCandidat.objects.update_or_create(
            candidat=candidat, reseau=reseau,
            defaults={'url': url},
        )
    else:
        LienCandidat.objects.filter(candidat=candidat, reseau=reseau).delete()


# ─── Génération asynchrone des artefacts ──────────────────────────────────────

def _schedule_artefacts_generation(cv_id, cv_data, request):
    """Lance la génération PDF + images dans un thread daemon.

    Dépose la requête de rendu et retourne immédiatement → la réponse HTTP du
    `sauvegarder_cv` n'attend plus les ~2-5s de Chromium/Playwright. Le thread :
      1. Ouvre sa propre connexion DB (via `close_old_connections` pour purger
         celle héritée du parent qui sera fermée à la fin de la requête).
      2. Re-fetch le CV (rebind ORM dans le thread).
      3. Appelle `generer_artefacts(request, cv_data=cv_data)`.
      4. Ferme la connexion DB.

    Erreurs : loggées, jamais propagées. La sauvegarde DB principale est déjà
    commit, donc un KO Playwright ne perd rien — le téléchargement régénère
    à la volée et un re-save retentera la génération.

    Limites : sous fort trafic, on lance N threads Chromium en parallèle
    (chacun ~150-300 MB RAM). Si ça devient un goulet, passer à une queue
    type Celery/RQ avec worker pool.
    """
    def _run():
        from django.db import close_old_connections
        try:
            close_old_connections()
            cv = CV.objects.filter(pk=cv_id).select_related('modele', 'contenu').first()
            if not cv:
                return
            cv.generer_artefacts(request, cv_data=cv_data)
        except Exception:
            logger.exception('Génération asynchrone artefacts CV %s a échoué', cv_id)
        finally:
            close_old_connections()

    threading.Thread(target=_run, name=f'cv-artefacts-{cv_id}', daemon=True).start()


@candidat_required
@require_POST
def sauvegarder_cv(request, template_id):
    """Sauvegarde un CV (modèle relationnel pur — sans JSON snapshot).

    Body JSON attendu :
        { "cv_id": <int|null>, "titre": "...", "cv": { ... données éditeur ... } }

    Pipeline :
      1. Mise à jour des champs identité/contact sur Candidat.
      2. Upsert de l'InformationPersonnelle (snapshot identité dans CVContenu).
      3. Synchronisation des rubriques relationnelles via `sync_rubriques_to_db`
         (delete+recreate côté Formation/Experience/Competence/Langue/CentreInteret
         /Projet/Benevolat). Les rubriques deviennent la source de vérité.
      4. Création / mise à jour du CV + CVContenu (FK info, M2M rubriques, toggles).

    Si `cv_id` est fourni et appartient au candidat, le CV est mis à jour ;
    sinon un nouveau CV est créé.
    """
    from .cv_initial import extract_rubriques_snapshot
    from .rubriques_sync import sync_rubriques_to_db

    try:
        payload = json.loads(request.body)
    except (json.JSONDecodeError, ValueError):
        return JsonResponse({'ok': False, 'message': 'Données invalides.'}, status=400)

    cv_data = payload.get('cv') or {}
    titre   = (payload.get('titre') or cv_data.get('titre') or '').strip()
    cv_id   = payload.get('cv_id')
    nom_cv  = (payload.get('nom_cv') or cv_data.get('nomCv') or '').strip()

    # Validation "Nom du CV" — obligatoire + unique par candidat.
    # On renvoie `field: 'nom_cv'` pour que l'éditeur cible l'erreur sur le bon
    # champ et bascule en §1 Informations personnelles.
    if not nom_cv:
        return JsonResponse(
            {'ok': False, 'field': 'nom_cv',
             'message': 'Veuillez saisir un nom pour ce CV.'},
            status=400,
        )
    if len(nom_cv) > 200:
        return JsonResponse(
            {'ok': False, 'field': 'nom_cv',
             'message': 'Nom du CV trop long (200 caractères max).'},
            status=400,
        )

    modele = ModeleCV.objects.filter(pk=template_id, actif=True).first()
    if not modele:
        return JsonResponse({'ok': False, 'message': 'Modèle introuvable.'}, status=404)

    candidat = request.candidat

    # Unicité par candidat (insensible à la casse pour éviter les confusions
    # « Mon CV » vs « mon cv »). Exclut le CV en cours d'édition.
    doublon_qs = CV.objects.filter(candidat=candidat, nomCv__iexact=nom_cv)
    if cv_id:
        doublon_qs = doublon_qs.exclude(pk=cv_id)
    if doublon_qs.exists():
        return JsonResponse(
            {'ok': False, 'field': 'nom_cv',
             'message': f"Vous avez déjà un CV nommé « {nom_cv} ». Choisissez un autre nom."},
            status=409,
        )

    try:
        with transaction.atomic():
            # 1. Identité directe sur Candidat
            _sync_candidat_identity(candidat, cv_data)

            # 2. Snapshot identité (FK depuis CVContenu)
            info = _upsert_information_personnelle(candidat, cv_data)

            # 3. Rubriques : sync **additif** sur le profil candidat.
            # On construit un supremum = items du CV + items du profil non
            # présents dans le CV. Le profil candidat (pool maître) ne perd
            # jamais de données — un CV simplifié n'érode plus la base.
            # Les items du CV viennent en premier dans `merged_data`, ce qui
            # garantit l'alignement par index avec `sync_result` pour le
            # calcul de `elementsMasques` ci-dessous.
            merged_data = _merge_cv_into_profile(candidat, cv_data)
            sync_result = sync_rubriques_to_db(candidat, merged_data)
            candidat.rubriques = extract_rubriques_snapshot(merged_data)
            candidat.save(update_fields=['rubriques'])

            # Carte de visibilité : pour chaque section, on collecte les IDs DB
            # des items marqués `visible:false` côté éditeur. `sync_result` est
            # une liste alignée index-par-index avec le JSON (None pour items
            # ignorés), ce qui permet de retrouver l'objet DB de chaque entrée.
            elements_masques = {}
            for section, json_key in [
                ('experiences', 'experiences'),
                ('formations',  'formations'),
                ('competences', 'competences'),
                ('langues',     'langues'),
                ('projets',     'projets'),
                ('benevs',      'benevs'),
            ]:
                items_json = cv_data.get(json_key) or []
                created    = sync_result.get(section) or []
                masked_ids = []
                for i, item in enumerate(items_json):
                    if not isinstance(item, dict):
                        continue
                    if item.get('visible') is False and i < len(created) and created[i] is not None:
                        masked_ids.append(created[i].pk)
                elements_masques[section] = masked_ids
            # Centres d'intérêt : indices (strings, pas d'objet → pas d'ID)
            elements_masques['interets'] = [
                i for i in (cv_data.get('interetsMasques') or [])
                if isinstance(i, int)
            ]

            # 4. CV + CVContenu
            cv_obj = None
            if cv_id:
                cv_obj = CV.objects.filter(pk=cv_id, candidat=candidat).first()
            if cv_obj is None:
                cv_obj = CV(candidat=candidat)

            cv_obj.modele = modele
            cv_obj.nomCv  = nom_cv[:200]
            cv_obj.titre  = titre[:200]
            cv_obj.profil = _txt(cv_data.get('profil'))

            # Photo : enregistre la nouvelle photo sur le CV (et sur le candidat
            # uniquement si sa photo de profil est vide). Voir _apply_cv_photo.
            # Doit être fait avant cv_obj.save() pour que le FileField soit
            # persisté dans la même transaction.
            _apply_cv_photo(cv_obj, candidat, cv_data.get('photo'))

            contenu = cv_obj.contenu or CVContenu.objects.create()
            contenu.informationPersonnelle = info
            contenu.showCertif      = bool(cv_data.get('showCertif',  False))
            contenu.showProjets     = bool(cv_data.get('showProjets', False))
            contenu.showBenev       = bool(cv_data.get('showBenev',   False))
            contenu.showRef         = bool(cv_data.get('showRef',     True))
            contenu.elementsMasques = elements_masques
            # Snapshot complet : ce dict est la source de vérité pour la
            # modification du CV. Le re-charger ressuscite le CV exactement
            # tel qu'il était au moment du save (rubriques, masquages, toggles,
            # identité). Indépendant des delete+recreate sur les rubriques
            # candidat → les CV passés ne sont plus cassés par un nouveau save.
            contenu.donneesSnapshot = _build_snapshot(cv_data, elements_masques)
            contenu.save()

            cv_obj.contenu = contenu
            cv_obj.save()

            # M2M : on lie les rubriques fraîchement (re)créées par sync_rubriques_to_db.
            contenu.formations.set(candidat.formations.all())
            contenu.experiences.set(candidat.experiencesProfessionnelles.all())
            contenu.competences.set(candidat.competences.all())
            contenu.langues.set(candidat.languesParlees.all())
            contenu.interets.set(candidat.centresInteret.all())
            contenu.projets.set(candidat.projets.all())
            contenu.benevolats.set(candidat.benevolats.all())
    except Exception as exc:
        return JsonResponse(
            {'ok': False, 'message': f"Erreur lors de la sauvegarde : {exc}"},
            status=500,
        )

    # 5. Génération des artefacts (PDF + images) en **arrière-plan**.
    # Playwright/Chromium prend ~2-5 secondes — on ne bloque plus la réponse.
    # Le thread daemon ré-ouvre une connexion DB propre et s'isole d'éventuelles
    # erreurs (best-effort : l'échec du rendu ne casse pas la sauvegarde DB).
    # Le PDF/images apparaîtront dans la liste « Mes CV » au prochain refresh.
    # Le téléchargement (`telecharger_cv`) régénère à la volée → indépendant.
    _schedule_artefacts_generation(cv_obj.pk, cv_data, request)

    return JsonResponse({
        'ok':              True,
        'cv_id':           cv_obj.pk,
        'message':         'CV sauvegardé.',
        # PDF de la précédente génération (peut être obsolète pour ce save).
        # `pdfPending=true` indique que la nouvelle version est en cours.
        'pdf_url':         cv_obj.cvPdf.url if cv_obj.cvPdf else '',
        'pdf_pending':     True,
    })
